"""
خادم Flask الرئيسي للنظام الخلفي
يوفر APIs للتفاعل مع نماذج الذكاء الاصطناعي وقاعدة البيانات
يدعم الآن استوديو الوكلاء المتقدم والواجهة الراقصة
"""

from flask import Flask, request, jsonify, send_file, Response, stream_with_context
from flask_cors import CORS
from dotenv import load_dotenv
import os
import io
import json
import time
import asyncio
import threading
from typing import Dict, Any, List, Optional, Generator
from datetime import datetime, timedelta

# استيراد الوحدات المحلية
from database import (
    init_db, save_project_data, get_project_data, create_new_project,
    get_all_projects, delete_project, update_project_stage
)
from llm_service import (
    call_llm, get_best_model_for_task, validate_api_keys,
    create_novel_analysis_prompt, create_idea_generation_prompt,
    create_blueprint_prompt, create_chapter_generation_prompt,
    create_text_refinement_prompt, create_consistency_check_prompt,
    create_suggestions_prompt, create_final_report_prompt
)
from adaptive_learning_service import get_adaptive_service
from database import (
    save_workflow_design, get_workflow_design, get_user_workflow_designs,
    delete_workflow_design, increment_workflow_usage
)

# استيراد مكونات نظام الوكلاء المتقدم
try:
    from agent_database import (
        init_agent_database, get_agent_by_id, get_all_agents, save_agent,
        update_agent, delete_agent, get_tool_by_id, get_all_tools,
        save_tool, update_tool, delete_tool, get_agent_messages,
        save_agent_message, get_agent_collaboration_session,
        create_agent_collaboration_session, update_collaboration_status
    )
    from advanced.AdvancedArbitrator import AdvancedArbitrator
    from advanced.AgentCollaboration import AgentCollaboration
    from personal_analytics_service import PersonalAnalyticsService
    
    # تهيئة مكونات نظام الوكلاء المتقدم
    arbitrator = AdvancedArbitrator()
    analytics_service = PersonalAnalyticsService()
    collaboration_system = AgentCollaboration()
    AGENT_STUDIO_ENABLED = True
except ImportError as e:
    print(f"تحذير: لم يتم تحميل نظام الوكلاء المتقدم: {e}")
    AGENT_STUDIO_ENABLED = False

# تحميل متغيرات البيئة
load_dotenv()

# إنشاء تطبيق Flask
app = Flask(__name__)

# تفعيل CORS لجميع المسارات (مهم للتطوير)
CORS(app, origins=["http://localhost:5173", "http://localhost:3000"])

# تسجيل البلو-برنت للأدوات المتخصصة
try:
    from specialized_tools_apis import specialized_tools_bp
    app.register_blueprint(specialized_tools_bp)
    print("✅ تم تحميل APIs الأدوات المتخصصة بنجاح")
except ImportError as e:
    print(f"⚠️ تحذير: لم يتم تحميل APIs الأدوات المتخصصة: {e}")

# تسجيل البلو-برنت لذاكرة السرد الحيّة
try:
    from living_narrative_apis import living_narrative_bp
    app.register_blueprint(living_narrative_bp)
    print("✅ تم تحميل APIs ذاكرة السرد الحيّة بنجاح")
except ImportError as e:
    print(f"⚠️ تحذير: لم يتم تحميل APIs ذاكرة السرد الحيّة: {e}")

# خدمة التعلم التكيفي
adaptive_service = get_adaptive_service()

def get_user_id_from_request(request) -> str:
    """استخراج معرف المستخدم من الطلب للتعلم التكيفي"""
    return adaptive_service.get_user_id(request)

# تهيئة قاعدة البيانات عند بدء التطبيق
with app.app_context():
    init_db()
    # تهيئة قاعدة بيانات الوكلاء إذا كان النظام المتقدم متاحاً
    if AGENT_STUDIO_ENABLED:
        init_agent_database()

# استيراد خدمة PDF المتقدمة
try:
    from advanced_pdf_service import get_pdf_service
    PDF_SERVICE_AVAILABLE = True
except ImportError:
    PDF_SERVICE_AVAILABLE = False

# معالجة قراءة الملفات
def read_file_content(file) -> tuple[str, str]:
    """قراءة محتوى الملف بناءً على نوعه - محسن مع دعم PDF متقدم"""
    try:
        file_type = file.mimetype
        file_content = ""
        
        if file_type == 'text/plain' or file.filename.endswith('.txt'):
            file_content = file.read().decode('utf-8')
            
        elif file_type == 'application/pdf' or file.filename.endswith('.pdf'):
            # استخدام خدمة PDF المتقدمة
            if PDF_SERVICE_AVAILABLE:
                try:
                    pdf_service = get_pdf_service()
                    pdf_data = file.read()
                    
                    # استخراج النص باستخدام الخدمة المتقدمة
                    extracted_text, error_message = pdf_service.extract_text_only(pdf_data)
                    
                    if error_message:
                        return "", f"خطأ في قراءة ملف PDF: {error_message}"
                    
                    file_content = extracted_text
                    
                except Exception as e:
                    # العودة للطريقة القديمة كـ fallback
                    return _fallback_pdf_read(file, str(e))
            else:
                # استخدام الطريقة القديمة
                return _fallback_pdf_read(file, "خدمة PDF المتقدمة غير متاحة")
                
        elif (file_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document' 
              or file.filename.endswith('.docx')):
            try:
                import mammoth
                # إعادة قراءة الملف للـ DOCX (في حال تم قراءته من قبل للـ PDF)
                file.seek(0)
                document_bytes = io.BytesIO(file.read())
                result = mammoth.extract_raw_text(document_bytes)
                file_content = result.value
            except Exception as e:
                return "", f"خطأ في قراءة ملف DOCX: {str(e)}"
        else:
            return "", f"صيغة الملف غير مدعومة: {file_type}"
        
        if not file_content.strip():
            return "", "الملف فارغ أو لا يحتوي على نص"
            
        return file_content, ""
        
    except Exception as e:
        return "", f"خطأ في قراءة الملف: {str(e)}"

def _fallback_pdf_read(file, initial_error: str) -> tuple[str, str]:
    """طريقة احتياطية لقراءة PDF باستخدام PyPDF2"""
    try:
        # إعادة تعيين موضع القراءة
        file.seek(0)
        
        from PyPDF2 import PdfReader
        reader = PdfReader(io.BytesIO(file.read()))
        
        file_content = ""
        for page in reader.pages:
            file_content += page.extract_text() + "\n"
        
        if not file_content.strip():
            return "", f"فشل في استخراج النص. {initial_error}"
        
        return file_content, ""
        
    except Exception as e:
        return "", f"خطأ في قراءة PDF (الطريقة الاحتياطية): {str(e)}. الخطأ الأولي: {initial_error}"

# ================== مسارات إدارة المشاريع ==================

@app.route('/api/health', methods=['GET'])
def health_check():
    """فحص صحة الخادم"""
    api_status = validate_api_keys()
    return jsonify({
        "status": "healthy",
        "message": "الخادم يعمل بشكل طبيعي",
        "api_keys_status": api_status,
        "timestamp": datetime.now().isoformat()
    })

@app.route('/api/projects', methods=['GET'])
def get_projects():
    """جلب قائمة بجميع المشاريع"""
    try:
        projects = get_all_projects()
        return jsonify({
            "success": True,
            "projects": projects,
            "count": len(projects)
        })
    except Exception as e:
        return jsonify({"error": f"خطأ في جلب المشاريع: {str(e)}"}), 500

@app.route('/api/projects', methods=['POST'])
def create_project():
    """إنشاء مشروع جديد"""
    try:
        data = request.json or {}
        title = data.get('title', 'مشروع جديد')
        description = data.get('description', 'وصف المشروع الجديد')
        
        project_id = create_new_project(title, description)
        
        if project_id:
            return jsonify({
                "success": True,
                "project_id": project_id,
                "message": "تم إنشاء المشروع بنجاح"
            }), 201
        else:
            return jsonify({"error": "فشل في إنشاء المشروع"}), 500
            
    except Exception as e:
        return jsonify({"error": f"خطأ في إنشاء المشروع: {str(e)}"}), 500

@app.route('/api/projects/<int:project_id>', methods=['GET'])
def get_project(project_id):
    """جلب بيانات مشروع محدد"""
    try:
        project = get_project_data(project_id)
        
        if project:
            return jsonify({
                "success": True,
                "project": project
            })
        else:
            return jsonify({"error": "المشروع غير موجود"}), 404
            
    except Exception as e:
        return jsonify({"error": f"خطأ في جلب المشروع: {str(e)}"}), 500

@app.route('/api/projects/<int:project_id>', methods=['DELETE'])
def delete_project_endpoint(project_id):
    """حذف مشروع"""
    try:
        success = delete_project(project_id)
        
        if success:
            return jsonify({
                "success": True,
                "message": "تم حذف المشروع بنجاح"
            })
        else:
            return jsonify({"error": "المشروع غير موجود"}), 404
            
    except Exception as e:
        return jsonify({"error": f"خطأ في حذف المشروع: {str(e)}"}), 500

@app.route('/api/projects/<int:project_id>/save', methods=['POST'])
def save_project_data_endpoint(project_id):
    """حفظ بيانات المشروع"""
    try:
        data = request.json
        
        if not data:
            return jsonify({"error": "لا توجد بيانات للحفظ"}), 400
            
        data_key = data.get('data_key')
        data_value = data.get('data_value')
        
        if not data_key or data_value is None:
            return jsonify({"error": "مفتاح البيانات أو القيمة مفقودة"}), 400
        
        success = save_project_data(project_id, data_key, data_value)
        
        if success:
            return jsonify({
                "success": True,
                "message": f"تم حفظ {data_key} للمشروع {project_id} بنجاح"
            })
        else:
            return jsonify({"error": "فشل في حفظ البيانات"}), 500
            
    except Exception as e:
        return jsonify({"error": f"خطأ في حفظ البيانات: {str(e)}"}), 500

@app.route('/api/projects/<int:project_id>/stage', methods=['PUT'])
def update_stage(project_id):
    """تحديث مرحلة المشروع"""
    try:
        data = request.json
        stage = data.get('stage')
        
        if not isinstance(stage, int) or stage < 1 or stage > 6:
            return jsonify({"error": "رقم المرحلة غير صحيح (1-6)"}), 400
        
        success = update_project_stage(project_id, stage)
        
        if success:
            return jsonify({
                "success": True,
                "message": f"تم تحديث المشروع إلى المرحلة {stage}"
            })
        else:
            return jsonify({"error": "فشل في تحديث المرحلة"}), 500
            
    except Exception as e:
        return jsonify({"error": f"خطأ في تحديث المرحلة: {str(e)}"}), 500

# ================== مسارات المراحل الست ==================

@app.route('/api/analyze-novel', methods=['POST'])
def analyze_novel():
    """المرحلة 1: تحليل الرواية المصدر"""
    try:
        if 'file' not in request.files:
            return jsonify({"error": "لم يتم تقديم ملف"}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({"error": "لم يتم اختيار ملف"}), 400
        
        # قراءة محتوى الملف
        file_content, error = read_file_content(file)
        if error:
            return jsonify({"error": error}), 400
        
        # تقليل حجم النص للتحليل (تجنب تجاوز حدود Token)
        words = file_content.split()
        max_words = 4000  # حد أقصى للكلمات
        if len(words) > max_words:
            truncated_content = " ".join(words[:max_words])
            print(f"⚠️ تم تقليص النص من {len(words)} إلى {max_words} كلمة")
        else:
            truncated_content = file_content
        
        # إنشاء prompt محسن للتحليل
        prompt_messages = create_novel_analysis_prompt(truncated_content)
        
        # استدعاء نموذج الذكاء الاصطناعي
        print("🔍 بدء تحليل النص...")
        analysis_result = call_llm(
            prompt_messages, 
            model=get_best_model_for_task("analysis"),
            max_tokens=4000,
            temperature=0.3  # درجة حرارة منخفضة للتحليل الدقيق
        )
        
        # التحقق من صحة النتيجة
        if "error" in analysis_result:
            return jsonify({
                "error": "فشل في التحليل بواسطة الذكاء الاصطناعي",
                "details": analysis_result
            }), 500
        
        # إضافة معلومات الملف
        analysis_result.update({
            "title": analysis_result.get("title", os.path.splitext(file.filename)[0]),
            "author": analysis_result.get("author", "غير محدد"),
            "content": file_content[:1000] + "..." if len(file_content) > 1000 else file_content,
            "uploadDate": datetime.now().isoformat(),
            "wordCount": len(words),
            "chapterCount": max(1, len(words) // 3000)  # تقدير تقريبي
        })
        
        return jsonify({
            "success": True,
            "analysis": analysis_result
        })
        
    except Exception as e:
        return jsonify({
            "error": f"خطأ في تحليل الملف: {str(e)}"
        }), 500

@app.route('/api/generate-ideas', methods=['POST'])
def generate_ideas():
    """المرحلة 2: توليد أفكار الروايات"""
    try:
        data = request.json or {}
        source_summary = data.get('source_analysis_summary', '')
        
        # إنشاء prompt محسن لتوليد الأفكار
        prompt_messages = create_idea_generation_prompt(source_summary)
        
        print("💡 بدء توليد الأفكار...")
        ideas_result = call_llm(
            prompt_messages,
            model=get_best_model_for_task("creative"),
            max_tokens=3000,
            temperature=0.8  # درجة حرارة عالية للإبداع
        )
        
        # التحقق من صحة النتيجة
        if "error" in ideas_result:
            return jsonify({
                "error": "فشل في توليد الأفكار",
                "details": ideas_result
            }), 500
        
        # استخراج الأفكار من النتيجة
        ideas = []
        if "ideas" in ideas_result:
            ideas = ideas_result["ideas"]
        elif isinstance(ideas_result, list):
            ideas = ideas_result
        else:
            return jsonify({
                "error": "تنسيق استجابة الذكاء الاصطناعي غير صحيح",
                "details": ideas_result
            }), 500
        
        # إضافة معرفات فريدة والحالة
        for i, idea in enumerate(ideas):
            idea["id"] = f"idea-{i+1}-{datetime.now().timestamp()}"
            idea["selected"] = False
        
        return jsonify({
            "success": True,
            "ideas": ideas,
            "count": len(ideas)
        })
        
    except Exception as e:
        return jsonify({
            "error": f"خطأ في توليد الأفكار: {str(e)}"
        }), 500

@app.route('/api/build-blueprint', methods=['POST'])
def build_blueprint():
    """المرحلة 3: بناء مخطط الرواية"""
    try:
        data = request.json
        selected_idea = data.get('selected_idea')
        
        if not selected_idea:
            return jsonify({"error": "لم يتم تحديد فكرة للرواية"}), 400
        
        # إنشاء prompt محسن لبناء المخطط
        prompt_messages = create_blueprint_prompt(selected_idea)
        
        print("🏗️ بدء بناء مخطط الرواية...")
        blueprint_result = call_llm(
            prompt_messages,
            model=get_best_model_for_task("blueprint"),
            max_tokens=6000,
            temperature=0.5
        )
        
        # التحقق من صحة النتيجة
        if "error" in blueprint_result:
            return jsonify({
                "error": "فشل في بناء المخطط",
                "details": blueprint_result
            }), 500
        
        return jsonify({
            "success": True,
            "blueprint": blueprint_result
        })
        
    except Exception as e:
        return jsonify({
            "error": f"خطأ في بناء المخطط: {str(e)}"
        }), 500

@app.route('/api/generate-chapter', methods=['POST'])
def generate_chapter():
    """المرحلة 4: توليد فصل من الرواية"""
    try:
        data = request.json
        chapter_blueprint = data.get('chapter_blueprint')
        novel_style_profile = data.get('novel_style_profile', {})
        previous_chapter_summary = data.get('previous_chapter_summary', '')
        
        if not chapter_blueprint:
            return jsonify({"error": "مخطط الفصل مفقود"}), 400
        
        # الحصول على معرف المستخدم للتخصيص الذكي
        user_id = get_user_id_from_request(request)
        
        # إنشاء prompt محسن لتوليد الفصل مع تفعيل أسلوب الجطلاوي
        prompt_messages = create_chapter_generation_prompt(
            chapter_blueprint, 
            novel_style_profile, 
            previous_chapter_summary,
            jattlaoui_style_enabled=True,  # تفعيل الأسلوب الجطلاوي المطور
            user_id=user_id  # للتخصيص الذكي
        )
        
        word_target = chapter_blueprint.get('wordTarget', 3000)
        
        print(f"✍️ بدء كتابة الفصل {chapter_blueprint.get('number', '؟')}...")
        chapter_result = call_llm(
            prompt_messages,
            model=get_best_model_for_task("creative"),
            max_tokens=word_target + 1000,  # مساحة إضافية للتقييم والتعليقات
            temperature=0.7
        )
        
        # التحقق من صحة النتيجة
        if "error" in chapter_result:
            return jsonify({
                "error": "فشل في توليد الفصل",
                "details": chapter_result
            }), 500
        
        # التأكد من وجود المحتوى المطلوب
        if "content" not in chapter_result:
            return jsonify({
                "error": "لم يتم توليد محتوى الفصل",
                "details": chapter_result
            }), 500
        
        return jsonify({
            "success": True,
            "chapter": chapter_result
        })
        
    except Exception as e:
        return jsonify({
            "error": f"خطأ في توليد الفصل: {str(e)}"
        }), 500

@app.route('/api/refine-text', methods=['POST'])
def refine_text():
    """المرحلة 5: تنقيح النص التفاعلي"""
    try:
        data = request.json
        text_selection = data.get('text_selection')
        operation = data.get('operation')  # 'rephrase', 'expand', 'summarize', 'improve'
        context = data.get('context', '')
        style_profile = data.get('novel_style_profile', {})
        
        if not text_selection or not operation:
            return jsonify({"error": "النص أو نوع العملية مفقود"}), 400
        
        # تحديد نوع العملية
        operations = {
            'rephrase': 'أعد صياغة هذا النص بأسلوب مختلف مع الحفاظ على المعنى',
            'expand': 'وسع هذا النص بتفاصيل إضافية ووصف أكثر ثراء',
            'summarize': 'لخص هذا النص بشكل مختصر ومركز',
            'improve': 'حسن هذا النص من ناحية الأسلوب والوضوح',
            'dialogue': 'حسن الحوار ليكون أكثر طبيعية وتعبيراً',
            'description': 'حسن الوصف ليكون أكثر حيوية وتفصيلاً'
        }
        
        operation_desc = operations.get(operation, 'حسن هذا النص')
        
        # الحصول على معرف المستخدم للتخصيص الذكي
        user_id = get_user_id_from_request(request)
        
        # إنشاء prompt محسن لتنقيح النص مع تفعيل أسلوب الجطلاوي
        prompt_messages = create_text_refinement_prompt(
            text_selection, operation, context, style_profile,
            jattlaoui_style_enabled=True,  # تفعيل الأسلوب الجطلاوي المطور
            user_id=user_id  # للتخصيص الذكي
        )
        
        print(f"✨ بدء تنقيح النص: {operation}")
        result = call_llm(
            prompt_messages,
            model=get_best_model_for_task("editing"),
            max_tokens=len(text_selection) * 2,  # ضعف طول النص الأصلي
            temperature=0.6,
            json_output=False  # نص عادي وليس JSON
        )
        
        # التحقق من صحة النتيجة
        if "error" in result:
            return jsonify({
                "error": "فشل في تنقيح النص",
                "details": result
            }), 500
        
        refined_text = result.get("content", str(result))
        
        return jsonify({
            "success": True,
            "refined_text": refined_text,
            "operation": operation
        })
        
    except Exception as e:
        return jsonify({
            "error": f"خطأ في تنقيح النص: {str(e)}"
        }), 500

@app.route('/api/check-consistency', methods=['POST'])
def check_consistency():
    """فحص اتساق الفصل مع الحبكة والشخصيات"""
    try:
        data = request.json
        chapter_content = data.get('chapter_content')
        character_profiles = data.get('character_profiles', [])
        plot_outline = data.get('plot_outline', {})
        
        if not chapter_content:
            return jsonify({"error": "محتوى الفصل مفقود"}), 400
        
        # إنشاء prompt محسن لفحص الاتساق
        prompt_messages = create_consistency_check_prompt(
            chapter_content, character_profiles, plot_outline
        )
        
        print("🔍 بدء فحص اتساق الفصل...")
        result = call_llm(
            prompt_messages,
            model=get_best_model_for_task("consistency"),
            max_tokens=2000,
            temperature=0.3
        )
        
        # التحقق من صحة النتيجة
        if "error" in result:
            return jsonify({
                "error": "فشل في فحص الاتساق",
                "details": result
            }), 500
        
        return jsonify({
            "success": True,
            "consistency_report": result
        })
        
    except Exception as e:
        return jsonify({
            "error": f"خطأ في فحص الاتساق: {str(e)}"
        }), 500

@app.route('/api/generate-final-report', methods=['POST'])
def generate_final_report():
    """المرحلة 6: إنشاء التقرير النهائي"""
    try:
        data = request.json
        chapters_summary = data.get('chapters_summary', [])
        blueprint_summary = data.get('blueprint_summary', {})
        source_analysis = data.get('source_analysis', {})
        
        # إنشاء prompt محسن للتقرير النهائي
        prompt_messages = create_final_report_prompt(
            chapters_summary, blueprint_summary, source_analysis
        )
        
        print("📊 بدء إنشاء التقرير النهائي...")
        result = call_llm(
            prompt_messages,
            model=get_best_model_for_task("analysis"),
            max_tokens=3000,
            temperature=0.4
        )
        
        # التحقق من صحة النتيجة
        if "error" in result:
            return jsonify({
                "error": "فشل في إنشاء التقرير",
                "details": result
            }), 500
        
        return jsonify({
            "success": True,
            "final_report": result
        })
        
    except Exception as e:
        return jsonify({
            "error": f"خطأ في إنشاء التقرير النهائي: {str(e)}"
        }), 500

@app.route('/api/export-novel', methods=['POST'])
def export_novel():
    """تصدير الرواية بصيغ مختلفة"""
    try:
        data = request.json
        chapters_content = data.get('chapters_content', [])
        export_format = data.get('format', 'txt')
        metadata = data.get('metadata', {})
        include_analysis = data.get('include_analysis', False)
        quality_report = data.get('quality_report', {})
        
        if not chapters_content:
            return jsonify({"error": "محتوى الفصول مفقود"}), 400
        
        # بناء النص الكامل
        full_text = ""
        
        # إضافة المعلومات الأساسية
        if metadata:
            full_text += f"عنوان الرواية: {metadata.get('title', 'غير محدد')}\n"
            full_text += f"الوصف: {metadata.get('description', 'غير محدد')}\n"
            full_text += f"تاريخ الإنشاء: {metadata.get('created_at', 'غير محدد')}\n"
            full_text += f"عدد الفصول: {len(chapters_content)}\n"
            full_text += "\n" + "="*50 + "\n\n"
        
        # إضافة الفصول
        for i, chapter in enumerate(chapters_content, 1):
            full_text += f"الفصل {i}\n"
            full_text += "-" * 20 + "\n\n"
            
            if isinstance(chapter, dict):
                full_text += chapter.get('content', str(chapter))
            else:
                full_text += str(chapter)
            
            full_text += "\n\n"
        
        # إضافة تقرير الجودة إذا كان مطلوباً
        if include_analysis and quality_report:
            full_text += "\n" + "="*50 + "\n"
            full_text += "تقرير الجودة والتحليل\n"
            full_text += "="*50 + "\n\n"
            
            if "overall_quality" in quality_report:
                full_text += f"الجودة الإجمالية: {quality_report['overall_quality']}%\n\n"
            
            if "strengths" in quality_report:
                full_text += "نقاط القوة:\n"
                for strength in quality_report['strengths']:
                    full_text += f"• {strength}\n"
                full_text += "\n"
            
            if "improvements" in quality_report:
                full_text += "نقاط التحسين:\n"
                for improvement in quality_report['improvements']:
                    full_text += f"• {improvement}\n"
                full_text += "\n"
        
        # إنشاء الملف وإرساله للتحميل
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        
        if export_format == 'txt':
            # إنشاء ملف نصي مؤقت
            filename = f"novel_{timestamp}.txt"
            file_buffer = io.BytesIO()
            file_buffer.write(full_text.encode('utf-8'))
            file_buffer.seek(0)
            
            return send_file(
                file_buffer,
                as_attachment=True,
                download_name=filename,
                mimetype='text/plain; charset=utf-8'
            )
        
        elif export_format == 'docx':
            try:
                from docx import Document
                from docx.shared import Inches
                
                # إنشاء مستند Word
                doc = Document()
                
                # إضافة العنوان
                if metadata and 'title' in metadata:
                    title = doc.add_heading(metadata['title'], 0)
                    title.alignment = 2  # مركز
                
                # إضافة معلومات الكتاب
                if metadata:
                    info_para = doc.add_paragraph()
                    if 'author' in metadata:
                        info_para.add_run(f"المؤلف: {metadata['author']}\n").bold = True
                    if 'description' in metadata:
                        info_para.add_run(f"الوصف: {metadata['description']}\n")
                    if 'total_words' in metadata:
                        info_para.add_run(f"عدد الكلمات: {metadata['total_words']:,}\n")
                    if 'chapters' in metadata:
                        info_para.add_run(f"عدد الفصول: {metadata['chapters']}\n")
                
                # إضافة خط فاصل
                doc.add_paragraph("=" * 50)
                doc.add_page_break()
                
                # إضافة الفصول
                for i, chapter in enumerate(chapters_content, 1):
                    chapter_heading = doc.add_heading(f"الفصل {i}", level=1)
                    chapter_heading.alignment = 1  # يسار
                    
                    # إضافة محتوى الفصل
                    chapter_para = doc.add_paragraph(chapter)
                    chapter_para.alignment = 3  # ضبط
                    
                    # فاصل بين الفصول
                    if i < len(chapters_content):
                        doc.add_page_break()
                
                # إضافة تقرير الجودة إذا كان مطلوباً
                if include_analysis and quality_report:
                    doc.add_page_break()
                    analysis_heading = doc.add_heading("تقرير الجودة والتحليل", level=1)
                    analysis_heading.alignment = 1
                    
                    if "overall_quality" in quality_report:
                        doc.add_paragraph(f"الجودة الإجمالية: {quality_report['overall_quality']}%")
                    
                    if "strengths" in quality_report:
                        strengths_para = doc.add_paragraph()
                        strengths_para.add_run("نقاط القوة:\n").bold = True
                        for strength in quality_report['strengths']:
                            strengths_para.add_run(f"• {strength}\n")
                    
                    if "improvements" in quality_report:
                        improvements_para = doc.add_paragraph()
                        improvements_para.add_run("نقاط التحسين:\n").bold = True
                        for improvement in quality_report['improvements']:
                            improvements_para.add_run(f"• {improvement}\n")
                
                # حفظ المستند في buffer
                filename = f"novel_{timestamp}.docx"
                file_buffer = io.BytesIO()
                doc.save(file_buffer)
                file_buffer.seek(0)
                
                return send_file(
                    file_buffer,
                    as_attachment=True,
                    download_name=filename,
                    mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                )
                
            except ImportError:
                return jsonify({"error": "مكتبة python-docx غير متاحة"}), 500
        
        elif export_format == 'json':
            export_data = {
                "metadata": metadata,
                "chapters": chapters_content,
                "quality_report": quality_report if include_analysis else None,
                "export_date": datetime.now().isoformat(),
                "total_words": sum(len(str(chapter).split()) for chapter in chapters_content)
            }
            
            filename = f"novel_{timestamp}.json"
            file_buffer = io.BytesIO()
            file_buffer.write(json.dumps(export_data, ensure_ascii=False, indent=2).encode('utf-8'))
            file_buffer.seek(0)
            
            return send_file(
                file_buffer,
                as_attachment=True,
                download_name=filename,
                mimetype='application/json; charset=utf-8'
            )
        
        else:
            return jsonify({"error": f"صيغة التصدير غير مدعومة: {export_format}"}), 400
        
    except Exception as e:
        return jsonify({
            "error": f"خطأ في تصدير الرواية: {str(e)}"
        }), 500

# ================== مسارات إضافية ==================

@app.route('/api/models', methods=['GET'])
def get_available_models():
    """جلب قائمة النماذج المتاحة"""
    from llm_service import AVAILABLE_MODELS, DEFAULT_MODEL
    
    return jsonify({
        "available_models": AVAILABLE_MODELS,
        "default_model": DEFAULT_MODEL,
        "api_status": validate_api_keys()
    })

@app.route('/api/test-llm', methods=['POST'])
def test_llm_connection():
    """اختبار الاتصال بنماذج الذكاء الاصطناعي"""
    try:
        test_messages = [
            {"role": "user", "content": "قل 'مرحباً! الاتصال يعمل بشكل طبيعي' بتنسيق JSON مع مفتاح 'message'."}
        ]
        
        result = call_llm(test_messages, task_type="creative", max_tokens=50)
        
        return jsonify({
            "success": "error" not in result,
            "result": result,
            "api_status": validate_api_keys()
        })
        
    except Exception as e:
        return jsonify({
            "success": False,
            "error": str(e),
            "api_status": validate_api_keys()
        })

# ================== معالج الأخطاء ==================

@app.errorhandler(404)
def not_found(error):
    return jsonify({"error": "المسار غير موجود"}), 404

@app.errorhandler(500)
def internal_error(error):
    return jsonify({"error": "خطأ داخلي في الخادم"}), 500

@app.errorhandler(413)
def request_entity_too_large(error):
    return jsonify({"error": "حجم الملف كبير جداً"}), 413

# ================== مسارات محرك سير العمل (Workflow Engine) ==================

from workflow_engine import (
    WorkflowEngine, WorkflowDefinition, WorkflowNode, NodeType, 
    ExecutionStatus, WorkflowTemplates, NodeFactory
)
import asyncio
from threading import Thread
import uuid
import time

# إنشاء محرك سير العمل
workflow_engine = WorkflowEngine()
active_workflow_executions = {}

@app.route('/api/workflows/templates', methods=['GET'])
def get_workflow_templates():
    """الحصول على قوالب سير العمل الجاهزة"""
    try:
        templates = [
            {
                "id": "complete_novel",
                "name": "سير عمل كامل لكتابة الرواية",
                "description": "من التحليل إلى الرواية النهائية",
                "complexity": "high",
                "estimated_time": "2-4 ساعات",
                "stages": ["تحليل", "أفكار", "مخطط", "فصول", "تحسين", "تقرير"]
            },
            {
                "id": "quick_ideas",
                "name": "توليد أفكار سريع",
                "description": "تحليل النص وتوليد أفكار فقط",
                "complexity": "low",
                "estimated_time": "15-30 دقيقة",
                "stages": ["تحليل", "أفكار"]
            }
        ]
        
        return jsonify({
            "success": True,
            "templates": templates
        })
        
    except Exception as e:
        return jsonify({"error": f"خطأ في جلب القوالب: {str(e)}"}), 500

@app.route('/api/workflows/create-from-template', methods=['POST'])
def create_workflow_from_template():
    """إنشاء سير عمل من قالب جاهز"""
    try:
        data = request.json or {}
        template_id = data.get('template_id')
        
        if not template_id:
            return jsonify({"error": "لم يتم تحديد معرف القالب"}), 400
        
        # إنشاء سير العمل حسب القالب
        if template_id == "complete_novel":
            workflow = WorkflowTemplates.create_complete_novel_workflow()
        elif template_id == "quick_ideas":
            workflow = WorkflowTemplates.create_quick_idea_workflow()
        else:
            return jsonify({"error": "قالب غير مدعوم"}), 400
        
        # تحويل سير العمل إلى تنسيق JSON للإرسال
        workflow_data = {
            "id": workflow.id,
            "name": workflow.name,
            "description": workflow.description,
            "nodes": [
                {
                    "id": node.id,
                    "type": node.type.value,
                    "name": node.name,
                    "description": node.description,
                    "config": node.config,
                    "inputs": node.inputs,
                    "outputs": node.outputs,
                    "position": node.position,
                    "status": node.status.value
                }
                for node in workflow.nodes
            ],
            "metadata": workflow.metadata
        }
        
        return jsonify({
            "success": True,
            "workflow": workflow_data
        })
        
    except Exception as e:
        return jsonify({"error": f"خطأ في إنشاء سير العمل: {str(e)}"}), 500

@app.route('/api/workflows/run', methods=['POST'])
def run_workflow():
    """تشغيل سير عمل"""
    try:
        data = request.json or {}
        
        # تحويل البيانات إلى كائن WorkflowDefinition
        workflow_data = data.get('workflow')
        if not workflow_data:
            return jsonify({"error": "لم يتم تقديم بيانات سير العمل"}), 400
        
        # إنشاء العقد
        nodes = []
        for node_data in workflow_data.get('nodes', []):
            node = WorkflowNode(
                id=node_data['id'],
                type=NodeType(node_data['type']),
                name=node_data['name'],
                description=node_data['description'],
                config=node_data.get('config', {}),
                inputs=node_data.get('inputs', []),
                outputs=node_data.get('outputs', []),
                position=node_data.get('position', {"x": 0, "y": 0}),
                status=ExecutionStatus.PENDING
            )
            nodes.append(node)
        
        # إنشاء كائن سير العمل
        workflow = WorkflowDefinition(
            id=workflow_data.get('id', str(uuid.uuid4())),
            name=workflow_data.get('name', 'سير عمل جديد'),
            description=workflow_data.get('description', ''),
            nodes=nodes,
            metadata=workflow_data.get('metadata', {}),
            created_at=datetime.now(),
            updated_at=datetime.now()
        )
        
        # البيانات الأولية (مثل ملف الرواية)
        initial_data = data.get('initial_data', {})
        
        # الحصول على معرف المستخدم للتعلم التكيفي
        user_id = get_user_id_from_request(request)
        
        # إنشاء معرف فريد للتنفيذ
        execution_id = f"exec_{int(time.time())}_{str(uuid.uuid4())[:8]}"
        
        # بدء التنفيذ في خيط منفصل لتجنب تعليق الواجهة
        def run_workflow_async():
            try:
                loop = asyncio.new_event_loop()
                asyncio.set_event_loop(loop)
                
                def progress_callback(progress):
                    active_workflow_executions[execution_id] = {
                        "progress": progress,
                        "last_update": datetime.now()
                    }
                
                result = loop.run_until_complete(
                    workflow_engine.execute_workflow(
                        workflow, 
                        initial_data, 
                        progress_callback,
                        user_id=user_id
                    )
                )
                
                # حفظ النتيجة النهائية
                active_workflow_executions[execution_id] = {
                    "result": result,
                    "completed": True,
                    "last_update": datetime.now()
                }
                
            except Exception as e:
                print(f"خطأ في تنفيذ سير العمل: {str(e)}")
                active_workflow_executions[execution_id] = {
                    "error": str(e),
                    "failed": True,
                    "last_update": datetime.now()
                }
        
        # بدء التنفيذ
        thread = Thread(target=run_workflow_async)
        thread.daemon = True
        thread.start()
        
        return jsonify({
            "success": True,
            "execution_id": execution_id,
            "message": "تم بدء تنفيذ سير العمل",
            "status": "running"
        })
        
    except Exception as e:
        return jsonify({"error": f"خطأ في تشغيل سير العمل: {str(e)}"}), 500

@app.route('/api/workflows/progress/<execution_id>', methods=['GET'])
def get_workflow_progress(execution_id):
    """الحصول على حالة تقدم سير العمل"""
    try:
        if execution_id not in active_workflow_executions:
            return jsonify({"error": "معرف التنفيذ غير موجود"}), 404
        
        execution_data = active_workflow_executions[execution_id]
        
        # التحقق من انتهاء التنفيذ
        if "result" in execution_data:
            return jsonify({
                "success": True,
                "completed": True,
                "result": execution_data["result"]
            })
        
        elif "error" in execution_data:
            return jsonify({
                "success": False,
                "failed": True,
                "error": execution_data["error"]
            })
        
        else:
            # التنفيذ قيد التقدم
            progress = execution_data.get("progress")
            if progress:
                return jsonify({
                    "success": True,
                    "running": True,
                    "progress": {
                        "total_nodes": progress.total_nodes,
                        "completed_nodes": progress.completed_nodes,
                        "failed_nodes": progress.failed_nodes,
                        "current_node": progress.current_node,
                        "progress_percentage": progress.progress_percentage,
                        "status": progress.status,
                        "logs": progress.logs[-10:] if progress.logs else []  # آخر 10 سجلات
                    }
                })
            else:
                return jsonify({
                    "success": True,
                    "running": True,
                    "progress": {
                        "status": "initializing",
                        "progress_percentage": 0
                    }
                })
        
    except Exception as e:
        return jsonify({"error": f"خطأ في جلب حالة التقدم: {str(e)}"}), 500

@app.route('/api/workflows/cancel/<execution_id>', methods=['POST'])
def cancel_workflow(execution_id):
    """إلغاء تنفيذ سير العمل"""
    try:
        if execution_id in active_workflow_executions:
            active_workflow_executions[execution_id] = {
                "cancelled": True,
                "last_update": datetime.now()
            }
            return jsonify({
                "success": True,
                "message": "تم إلغاء تنفيذ سير العمل"
            })
        else:
            return jsonify({"error": "معرف التنفيذ غير موجود"}), 404
        
    except Exception as e:
        return jsonify({"error": f"خطأ في إلغاء سير العمل: {str(e)}"}), 500

@app.route('/api/workflows/save', methods=['POST'])
def save_workflow():
    """حفظ تصميم سير العمل في قاعدة البيانات"""
    try:
        data = request.json or {}
        workflow_design = data.get('workflow')
        
        if not workflow_design:
            return jsonify({"error": "لم يتم تقديم بيانات سير العمل"}), 400
        
        # الحصول على معرف المستخدم
        user_id = get_user_id_from_request(request)
        
        # استخراج البيانات من التصميم
        name = workflow_design.get('name', 'سير عمل جديد')
        description = workflow_design.get('description', '')
        is_template = data.get('is_template', False)
        is_public = data.get('is_public', False)
        tags = data.get('tags', [])
        complexity_level = data.get('complexity_level', 'medium')
        
        # تقدير مدة التنفيذ بناءً على عدد العقد
        node_count = len(workflow_design.get('nodes', []))
        estimated_duration = node_count * 15  # 15 دقيقة لكل عقدة تقريباً
        
        # حفظ في قاعدة البيانات
        workflow_id = save_workflow_design(
            user_id=user_id,
            name=name,
            description=description,
            workflow_json=workflow_design,
            is_template=is_template,
            is_public=is_public,
            tags=tags,
            complexity_level=complexity_level,
            estimated_duration_minutes=estimated_duration
        )
        
        if workflow_id:
            return jsonify({
                "success": True,
                "workflow_id": workflow_id,
                "message": "تم حفظ سير العمل بنجاح في قاعدة البيانات"
            })
        else:
            return jsonify({"error": "فشل في حفظ سير العمل"}), 500
        
    except Exception as e:
        return jsonify({"error": f"خطأ في حفظ سير العمل: {str(e)}"}), 500

@app.route('/api/workflows/load/<int:workflow_id>', methods=['GET'])
def load_workflow(workflow_id):
    """تحميل تصميم سير العمل المحفوظ من قاعدة البيانات"""
    try:
        # تحميل من قاعدة البيانات
        workflow_design = get_workflow_design(workflow_id)
        
        if not workflow_design:
            return jsonify({"error": "سير العمل غير موجود"}), 404
        
        # زيادة عداد الاستخدام
        increment_workflow_usage(workflow_id)
        
        return jsonify({
            "success": True,
            "workflow": workflow_design['workflow_json'],
            "metadata": {
                "id": workflow_design['id'],
                "name": workflow_design['name'],
                "description": workflow_design['description'],
                "created_at": workflow_design['created_at'],
                "updated_at": workflow_design['updated_at'],
                "usage_count": workflow_design['usage_count'],
                "tags": workflow_design['tags_json'],
                "complexity_level": workflow_design['complexity_level'],
                "estimated_duration_minutes": workflow_design['estimated_duration_minutes']
            }
        })
        
    except Exception as e:
        return jsonify({"error": f"خطأ في تحميل سير العمل: {str(e)}"}), 500

@app.route('/api/workflows/list', methods=['GET'])
def list_saved_workflows():
    """الحصول على قائمة بسير العمل المحفوظة من قاعدة البيانات"""
    try:
        # الحصول على معرف المستخدم
        user_id = get_user_id_from_request(request)
        
        # جلب تصميمات سير العمل للمستخدم (مع القوالب العامة)
        workflows_data = get_user_workflow_designs(user_id, include_public=True)
        
        # تحويل لتنسيق API
        workflows = []
        for workflow in workflows_data:
            workflows.append({
                "id": workflow['id'],
                "name": workflow['name'],
                "description": workflow['description'],
                "created_at": workflow['created_at'],
                "updated_at": workflow['updated_at'],
                "node_count": len(workflow['workflow_json'].get('nodes', [])),
                "usage_count": workflow['usage_count'],
                "tags": workflow['tags_json'],
                "complexity_level": workflow['complexity_level'],
                "estimated_duration_minutes": workflow['estimated_duration_minutes'],
                "is_template": workflow['is_template'],
                "is_public": workflow['is_public'],
                "is_mine": workflow['user_identifier'] == user_id
            })
        
        return jsonify({
            "success": True,
            "workflows": workflows,
            "total_count": len(workflows)
        })
        
    except Exception as e:
        return jsonify({"error": f"خطأ في جلب قائمة سير العمل: {str(e)}"}), 500

@app.route('/api/workflows/delete/<int:workflow_id>', methods=['DELETE'])
def delete_workflow_api(workflow_id):
    """حذف تصميم سير العمل"""
    try:
        # الحصول على معرف المستخدم
        user_id = get_user_id_from_request(request)
        
        # حذف التصميم (فقط إذا كان المستخدم هو المالك)
        success = delete_workflow_design(workflow_id, user_id)
        
        if success:
            return jsonify({
                "success": True,
                "message": "تم حذف سير العمل بنجاح"
            })
        else:
            return jsonify({"error": "لم يتم العثور على سير العمل أو ليس لديك صلاحية لحذفه"}), 404
        
    except Exception as e:
        return jsonify({"error": f"خطأ في حذف سير العمل: {str(e)}"}), 500

# ================== APIs التعلم التكيفي وتقييم المحتوى ==================

@app.route('/api/content/rate', methods=['POST'])
def rate_content():
    """تقييم المحتوى المُولد لأغراض التعلم التكيفي"""
    try:
        data = request.json or {}
        
        # الحصول على معرف المستخدم
        user_id = get_user_id_from_request(request)
        
        # استخراج بيانات التقييم
        content_type = data.get('content_type', 'general')
        content_text = data.get('content', '')
        rating_value = data.get('rating', 3)
        specific_feedback = data.get('feedback', {})
        project_id = data.get('project_id')
        
        # التحقق من صحة التقييم
        if not (1 <= rating_value <= 5):
            return jsonify({"error": "التقييم يجب أن يكون بين 1 و 5"}), 400
        
        # حفظ التقييم
        success = adaptive_service.save_user_rating(
            user_id=user_id,
            content_type=content_type,
            content=content_text,
            rating=rating_value,
            specific_feedback=specific_feedback,
            project_id=project_id
        )
        
        if success:
            # تحديث التعلم التكيفي
            adaptive_service.learn_from_interactions(user_id)
            
            return jsonify({
                "success": True,
                "message": "تم حفظ التقييم وتحديث التفضيلات بنجاح"
            })
        else:
            return jsonify({"error": "فشل في حفظ التقييم"}), 500
        
    except Exception as e:
        return jsonify({"error": f"خطأ في حفظ التقييم: {str(e)}"}), 500

@app.route('/api/user/profile', methods=['GET'])
def get_user_profile():
    """الحصول على ملف تعريف المستخدم والتفضيلات"""
    try:
        # الحصول على معرف المستخدم
        user_id = get_user_id_from_request(request)
        
        # تهيئة ملف التعريف
        profile = adaptive_service.initialize_user_profile(user_id)
        
        if profile:
            # إخفاء المعلومات الحساسة
            safe_profile = {
                'user_id': user_id,
                'profile_name': profile.get('profile_name', 'الملف الشخصي'),
                'style_preferences': profile.get('style_preferences', {}),
                'writing_habits': profile.get('writing_habits', {}),
                'jattlaoui_adaptation_level': profile.get('jattlaoui_adaptation_level', 0.5),
                'preferred_vocabulary_complexity': profile.get('preferred_vocabulary_complexity', 0.7),
                'preferred_sentence_length': profile.get('preferred_sentence_length', 0.6),
                'preferred_cultural_depth': profile.get('preferred_cultural_depth', 0.8),
                'created_at': profile.get('created_at'),
                'updated_at': profile.get('updated_at')
            }
            
            return jsonify({
                "success": True,
                "profile": safe_profile
            })
        else:
            return jsonify({"error": "فشل في جلب ملف التعريف"}), 500
        
    except Exception as e:
        return jsonify({"error": f"خطأ في جلب ملف التعريف: {str(e)}"}), 500

@app.route('/api/user/preferences', methods=['PUT'])
def update_user_preferences():
    """تحديث تفضيلات المستخدم"""
    try:
        data = request.json or {}
        
        # الحصول على معرف المستخدم
        user_id = get_user_id_from_request(request)
        
        # استخراج التحديثات
        updates = {}
        
        if 'style_preferences' in data:
            updates['style_preferences_json'] = data['style_preferences']
        
        if 'writing_habits' in data:
            updates['writing_habits_json'] = data['writing_habits']
        
        if 'jattlaoui_adaptation_level' in data:
            level = float(data['jattlaoui_adaptation_level'])
            if 0 <= level <= 1:
                updates['jattlaoui_adaptation_level'] = level
        
        if 'preferred_vocabulary_complexity' in data:
            level = float(data['preferred_vocabulary_complexity'])
            if 0 <= level <= 1:
                updates['preferred_vocabulary_complexity'] = level
        
        if 'preferred_sentence_length' in data:
            level = float(data['preferred_sentence_length'])
            if 0 <= level <= 1:
                updates['preferred_sentence_length'] = level
        
        if 'preferred_cultural_depth' in data:
            level = float(data['preferred_cultural_depth'])
            if 0 <= level <= 1:
                updates['preferred_cultural_depth'] = level
        
        # تطبيق التحديثات
        if updates:
            from database import update_writer_profile
            success = update_writer_profile(user_id, updates)
            
            if success:
                return jsonify({
                    "success": True,
                    "message": "تم تحديث التفضيلات بنجاح"
                })
            else:
                return jsonify({"error": "فشل في تحديث التفضيلات"}), 500
        else:
            return jsonify({"error": "لا توجد تحديثات صالحة"}), 400
        
    except Exception as e:
        return jsonify({"error": f"خطأ في تحديث التفضيلات: {str(e)}"}), 500

@app.route('/api/user/stats', methods=['GET'])
def get_user_stats():
    """الحصول على إحصائيات المستخدم"""
    try:
        # الحصول على معرف المستخدم
        user_id = get_user_id_from_request(request)
        
        # تحليل التفضيلات
        from database import analyze_user_preferences
        preferences_analysis = analyze_user_preferences(user_id)
        
        # جلب المشاريع والتصميمات
        from database import get_all_projects, get_user_workflow_designs
        user_projects = get_all_projects()
        user_workflows = get_user_workflow_designs(user_id, include_public=False)
        
        stats = {
            'total_projects': len(user_projects) if user_projects else 0,
            'total_workflows': len(user_workflows) if user_workflows else 0,
            'total_ratings': sum(
                analysis.get('rating_count', 0) 
                for analysis in preferences_analysis.get('rating_analysis', {}).values()
            ),
            'average_satisfaction': 0,  # يحتاج حساب من التقييمات
            'learning_progress': preferences_analysis.get('modification_patterns', []),
            'activity_summary': preferences_analysis.get('interaction_analysis', {}),
            'last_activity': preferences_analysis.get('updated_at')
        }
        
        # حساب متوسط الرضا
        rating_analysis = preferences_analysis.get('rating_analysis', {})
        if rating_analysis:
            total_weighted_rating = 0
            total_count = 0
            for content_type, analysis in rating_analysis.items():
                avg_rating = analysis.get('average_rating', 3)
                count = analysis.get('rating_count', 0)
                total_weighted_rating += avg_rating * count
                total_count += count
            
            if total_count > 0:
                stats['average_satisfaction'] = round(total_weighted_rating / total_count, 2)
        
        return jsonify({
            "success": True,
            "stats": stats
        })
        
    except Exception as e:
        return jsonify({"error": f"خطأ في جلب الإحصائيات: {str(e)}"}), 500

# ================== مسارات استوديو الوكلاء ==================

@app.route('/api/agents', methods=['GET'])
def get_agents():
    """الحصول على قائمة الوكلاء المتاحين"""
    if not AGENT_STUDIO_ENABLED:
        return jsonify({"error": "نظام الوكلاء المتقدم غير متاح"}), 404
    
    try:
        agents = get_all_agents()
        return jsonify({
            "success": True,
            "agents": agents
        })
    except Exception as e:
        return jsonify({"error": f"خطأ في استرجاع الوكلاء: {str(e)}"}), 500

@app.route('/api/agents/<agent_id>', methods=['GET'])
def get_agent(agent_id):
    """الحصول على تفاصيل وكيل محدد"""
    if not AGENT_STUDIO_ENABLED:
        return jsonify({"error": "نظام الوكلاء المتقدم غير متاح"}), 404
    
    try:
        agent = get_agent_by_id(agent_id)
        if not agent:
            return jsonify({"error": "الوكيل غير موجود"}), 404
        
        return jsonify({
            "success": True,
            "agent": agent
        })
    except Exception as e:
        return jsonify({"error": f"خطأ في استرجاع الوكيل: {str(e)}"}), 500

@app.route('/api/agents', methods=['POST'])
def create_agent():
    """إنشاء وكيل جديد"""
    if not AGENT_STUDIO_ENABLED:
        return jsonify({"error": "نظام الوكلاء المتقدم غير متاح"}), 404
    
    try:
        data = request.get_json()
        
        # التحقق من البيانات المطلوبة
        required_fields = ['name', 'type', 'description']
        for field in required_fields:
            if field not in data:
                return jsonify({"error": f"الحقل {field} مطلوب"}), 400
        
        # حفظ الوكيل الجديد
        agent_id = save_agent(data)
        return jsonify({
            "success": True,
            "agent_id": agent_id,
            "message": "تم إنشاء الوكيل بنجاح"
        })
    except Exception as e:
        return jsonify({"error": f"خطأ في إنشاء الوكيل: {str(e)}"}), 500

@app.route('/api/agents/<agent_id>', methods=['PUT'])
def update_agent_route(agent_id):
    """تحديث معلومات وكيل"""
    if not AGENT_STUDIO_ENABLED:
        return jsonify({"error": "نظام الوكلاء المتقدم غير متاح"}), 404
    
    try:
        data = request.get_json()
        
        # التحقق من وجود الوكيل
        agent = get_agent_by_id(agent_id)
        if not agent:
            return jsonify({"error": "الوكيل غير موجود"}), 404
        
        # تحديث الوكيل
        success = update_agent(agent_id, data)
        if not success:
            return jsonify({"error": "فشل تحديث الوكيل"}), 500
        
        return jsonify({
            "success": True,
            "message": "تم تحديث الوكيل بنجاح"
        })
    except Exception as e:
        return jsonify({"error": f"خطأ في تحديث الوكيل: {str(e)}"}), 500

@app.route('/api/agents/<agent_id>', methods=['DELETE'])
def delete_agent_route(agent_id):
    """حذف وكيل"""
    if not AGENT_STUDIO_ENABLED:
        return jsonify({"error": "نظام الوكلاء المتقدم غير متاح"}), 404
    
    try:
        # التحقق من وجود الوكيل
        agent = get_agent_by_id(agent_id)
        if not agent:
            return jsonify({"error": "الوكيل غير موجود"}), 404
        
        # حذف الوكيل
        success = delete_agent(agent_id)
        if not success:
            return jsonify({"error": "فشل حذف الوكيل"}), 500
        
        return jsonify({
            "success": True,
            "message": "تم حذف الوكيل بنجاح"
        })
    except Exception as e:
        return jsonify({"error": f"خطأ في حذف الوكيل: {str(e)}"}), 500

@app.route('/api/tools', methods=['GET'])
def get_tools():
    """الحصول على قائمة الأدوات المتاحة"""
    if not AGENT_STUDIO_ENABLED:
        return jsonify({"error": "نظام الوكلاء المتقدم غير متاح"}), 404
    
    try:
        tools = get_all_tools()
        return jsonify({
            "success": True,
            "tools": tools
        })
    except Exception as e:
        return jsonify({"error": f"خطأ في استرجاع الأدوات: {str(e)}"}), 500

@app.route('/api/tools/<tool_id>', methods=['GET'])
def get_tool(tool_id):
    """الحصول على تفاصيل أداة محددة"""
    if not AGENT_STUDIO_ENABLED:
        return jsonify({"error": "نظام الوكلاء المتقدم غير متاح"}), 404
    
    try:
        tool = get_tool_by_id(tool_id)
        if not tool:
            return jsonify({"error": "الأداة غير موجودة"}), 404
        
        return jsonify({
            "success": True,
            "tool": tool
        })
    except Exception as e:
        return jsonify({"error": f"خطأ في استرجاع الأداة: {str(e)}"}), 500

@app.route('/api/tools', methods=['POST'])
def create_tool():
    """إنشاء أداة جديدة"""
    if not AGENT_STUDIO_ENABLED:
        return jsonify({"error": "نظام الوكلاء المتقدم غير متاح"}), 404
    
    try:
        data = request.get_json()
        
        # التحقق من البيانات المطلوبة
        required_fields = ['name', 'category', 'description', 'function_name']
        for field in required_fields:
            if field not in data:
                return jsonify({"error": f"الحقل {field} مطلوب"}), 400
        
        # حفظ الأداة الجديدة
        tool_id = save_tool(data)
        return jsonify({
            "success": True,
            "tool_id": tool_id,
            "message": "تم إنشاء الأداة بنجاح"
        })
    except Exception as e:
        return jsonify({"error": f"خطأ في إنشاء الأداة: {str(e)}"}), 500

@app.route('/api/agents/<agent_id>/messages', methods=['GET'])
def get_agent_messages_route(agent_id):
    """الحصول على رسائل وكيل معين"""
    if not AGENT_STUDIO_ENABLED:
        return jsonify({"error": "نظام الوكلاء المتقدم غير متاح"}), 404
    
    try:
        # التحقق من وجود الوكيل
        agent = get_agent_by_id(agent_id)
        if not agent:
            return jsonify({"error": "الوكيل غير موجود"}), 404
        
        limit = request.args.get('limit', default=20, type=int)
        offset = request.args.get('offset', default=0, type=int)
        
        messages = get_agent_messages(agent_id, limit, offset)
        return jsonify({
            "success": True,
            "messages": messages
        })
    except Exception as e:
        return jsonify({"error": f"خطأ في استرجاع الرسائل: {str(e)}"}), 500

@app.route('/api/agents/messages', methods=['POST'])
def send_agent_message():
    """إرسال رسالة إلى وكيل"""
    if not AGENT_STUDIO_ENABLED:
        return jsonify({"error": "نظام الوكلاء المتقدم غير متاح"}), 404
    
    try:
        data = request.get_json()
        
        # التحقق من البيانات المطلوبة
        required_fields = ['from_agent_id', 'to_agent_id', 'content']
        for field in required_fields:
            if field not in data:
                return jsonify({"error": f"الحقل {field} مطلوب"}), 400
        
        # حفظ الرسالة
        message_id = save_agent_message(data)
        
        # معالجة الرسالة بواسطة نظام التعاون إذا كان ذلك مطلوبًا
        if data.get('collaboration_session_id'):
            collaboration_system.process_message(data)
        
        return jsonify({
            "success": True,
            "message_id": message_id,
            "message": "تم إرسال الرسالة بنجاح"
        })
    except Exception as e:
        return jsonify({"error": f"خطأ في إرسال الرسالة: {str(e)}"}), 500

@app.route('/api/agent-collaborations', methods=['POST'])
def create_collaboration():
    """إنشاء جلسة تعاون جديدة بين الوكلاء"""
    if not AGENT_STUDIO_ENABLED:
        return jsonify({"error": "نظام الوكلاء المتقدم غير متاح"}), 404
    
    try:
        data = request.get_json()
        
        # التحقق من البيانات المطلوبة
        required_fields = ['name', 'agent_ids', 'purpose']
        for field in required_fields:
            if field not in data:
                return jsonify({"error": f"الحقل {field} مطلوب"}), 400
        
        # إنشاء جلسة تعاون
        session_id = create_agent_collaboration_session(data)
        
        # بدء الجلسة
        threading.Thread(
            target=collaboration_system.start_session,
            args=(session_id,)
        ).start()
        
        return jsonify({
            "success": True,
            "session_id": session_id,
            "message": "تم إنشاء جلسة التعاون بنجاح"
        })
    except Exception as e:
        return jsonify({"error": f"خطأ في إنشاء جلسة التعاون: {str(e)}"}), 500

@app.route('/api/agent-collaborations/<session_id>/messages', methods=['GET'])
def get_collaboration_messages(session_id):
    """الحصول على رسائل جلسة تعاون"""
    if not AGENT_STUDIO_ENABLED:
        return jsonify({"error": "نظام الوكلاء المتقدم غير متاح"}), 404
    
    try:
        # التحقق من وجود الجلسة
        session = get_agent_collaboration_session(session_id)
        if not session:
            return jsonify({"error": "جلسة التعاون غير موجودة"}), 404
        
        messages = collaboration_system.get_session_messages(session_id)
        return jsonify({
            "success": True,
            "messages": messages
        })
    except Exception as e:
        return jsonify({"error": f"خطأ في استرجاع رسائل الجلسة: {str(e)}"}), 500

@app.route('/api/agent-collaborations/<session_id>/brainstorm', methods=['POST'])
def start_brainstorming(session_id):
    """بدء عصف ذهني في جلسة تعاون"""
    if not AGENT_STUDIO_ENABLED:
        return jsonify({"error": "نظام الوكلاء المتقدم غير متاح"}), 404
    
    try:
        data = request.get_json()
        
        # التحقق من وجود الجلسة
        session = get_agent_collaboration_session(session_id)
        if not session:
            return jsonify({"error": "جلسة التعاون غير موجودة"}), 404
        
        # بدء العصف الذهني
        brainstorm_id = collaboration_system.start_brainstorm(
            session_id, 
            data.get('topic'),
            data.get('duration', 60),
            data.get('rounds', 3)
        )
        
        return jsonify({
            "success": True,
            "brainstorm_id": brainstorm_id,
            "message": "تم بدء العصف الذهني بنجاح"
        })
    except Exception as e:
        return jsonify({"error": f"خطأ في بدء العصف الذهني: {str(e)}"}), 500

@app.route('/api/contents/evaluate', methods=['POST'])
def evaluate_content():
    """تقييم محتوى باستخدام المحكم المتقدم"""
    if not AGENT_STUDIO_ENABLED:
        return jsonify({"error": "نظام الوكلاء المتقدم غير متاح"}), 404
    
    try:
        data = request.get_json()
        
        # التحقق من البيانات المطلوبة
        if 'content' not in data:
            return jsonify({"error": "الحقل content مطلوب"}), 400
        
        # الحصول على معايير التقييم
        criteria = data.get('criteria', {
            'style': True,
            'grammar': True,
            'coherence': True,
            'creativity': True,
            'overall': True
        })
        
        # تقييم المحتوى
        evaluation = arbitrator.evaluate_content(data['content'], criteria)
        
        return jsonify({
            "success": True,
            "evaluation": evaluation
        })
    except Exception as e:
        return jsonify({"error": f"خطأ في تقييم المحتوى: {str(e)}"}), 500

@app.route('/api/contents/refine', methods=['POST'])
def refine_content():
    """تحسين محتوى باستخدام المحكم المتقدم"""
    if not AGENT_STUDIO_ENABLED:
        return jsonify({"error": "نظام الوكلاء المتقدم غير متاح"}), 404
    
    try:
        data = request.get_json()
        
        # التحقق من البيانات المطلوبة
        if 'content' not in data:
            return jsonify({"error": "الحقل content مطلوب"}), 400
        
        # الحصول على معايير التحسين
        criteria = data.get('criteria', {
            'style': True,
            'grammar': True,
            'coherence': True,
            'creativity': False
        })
        
        # تحسين المحتوى
        improved_content = arbitrator.refine_content(data['content'], criteria)
        
        return jsonify({
            "success": True,
            "improved_content": improved_content
        })
    except Exception as e:
        return jsonify({"error": f"خطأ في تحسين المحتوى: {str(e)}"}), 500

# ===============================
# APIs الرحلة الموحدة
# ===============================

@app.route('/api/unified-journey/templates', methods=['GET'])
def get_journey_templates():
    """جلب قوالب الرحلة الموحدة المتاحة"""
    try:
        # قوالب افتراضية للرحلة الموحدة
        default_templates = [
            {
                'id': 'complete_analysis',
                'name': 'رحلة التحليل الشاملة',
                'description': 'رحلة كاملة تبدأ من تحليل رواية مصدر',
                'mode': 'guided',
                'stages': [1, 2, 3, 4, 5, 6],
                'estimated_time': '2-3 ساعات',
                'difficulty': 'متوسط',
                'category': 'analysis_based'
            },
            {
                'id': 'creative_journey',
                'name': 'رحلة الإبداع الحر',
                'description': 'رحلة إبداعية تبدأ من توليد الأفكار',
                'mode': 'guided',
                'stages': [2, 3, 4, 5, 6],
                'estimated_time': '1.5-2 ساعة',
                'difficulty': 'سهل',
                'category': 'creative'
            },
            {
                'id': 'custom_workflow',
                'name': 'سير عمل مخصص',
                'description': 'تحكم كامل في بناء سير العمل',
                'mode': 'custom',
                'stages': 'customizable',
                'estimated_time': 'متغير',
                'difficulty': 'متقدم',
                'category': 'custom'
            },
            {
                'id': 'hybrid_approach',
                'name': 'النهج المختلط',
                'description': 'توازن بين التوجيه والمرونة',
                'mode': 'hybrid',
                'stages': [1, 2, 3, 4, 5, 6],
                'estimated_time': '2-3 ساعات',
                'difficulty': 'متوسط',
                'category': 'hybrid'
            }
        ]
        
        # جلب القوالب المخصصة من قاعدة البيانات
        user_id = get_user_id_from_request(request)
        user_templates = get_user_workflow_designs(user_id)
        
        # دمج القوالب الافتراضية والمخصصة
        all_templates = default_templates + [
            {
                'id': f"user_{template['id']}",
                'name': template['name'],
                'description': template['description'],
                'mode': 'custom',
                'stages': 'customizable',
                'estimated_time': 'متغير',
                'difficulty': 'مخصص',
                'category': 'user_created',
                'workflow_data': json.loads(template['workflow_data_json'])
            }
            for template in user_templates
            if template.get('is_template', False)
        ]
        
        return jsonify({
            'success': True,
            'templates': all_templates
        })
        
    except Exception as e:
        return jsonify({'error': f'خطأ في جلب القوالب: {str(e)}'}), 500

@app.route('/api/unified-journey/create-from-template', methods=['POST'])
def create_journey_from_template():
    """إنشاء رحلة موحدة من قالب"""
    try:
        data = request.get_json()
        template_id = data.get('template_id')
        journey_mode = data.get('journey_mode', 'guided')
        start_choice = data.get('start_choice', 'analyze')
        project_name = data.get('project_name', 'رحلة جديدة')
        
        if not template_id:
            return jsonify({'error': 'معرف القالب مطلوب'}), 400
        
        user_id = get_user_id_from_request(request)
        
        # إنشاء مشروع جديد للرحلة
        project_id = create_new_project(project_name, f'رحلة موحدة - {journey_mode}')
        
        # إنشاء سير العمل من القالب
        journey_workflow = create_journey_workflow(
            template_id, journey_mode, start_choice, project_id
        )
        
        # حفظ سير العمل
        workflow_id = save_workflow_design(
            user_id,
            f"رحلة_{project_name}",
            f"سير عمل للرحلة الموحدة - {journey_mode}",
            journey_workflow,
            is_template=False,
            template_category='unified_journey'
        )
        
        return jsonify({
            'success': True,
            'project_id': project_id,
            'workflow_id': workflow_id,
            'journey_data': journey_workflow
        })
        
    except Exception as e:
        return jsonify({'error': f'خطأ في إنشاء الرحلة: {str(e)}'}), 500

@app.route('/api/unified-journey/save-progress', methods=['POST'])
def save_journey_progress():
    """حفظ تقدم الرحلة الموحدة"""
    try:
        data = request.get_json()
        user_id = get_user_id_from_request(request)
        
        template_id = data.get('template_id')
        current_stage = data.get('current_stage', 0)
        completed_stages = data.get('completed_stages', [])
        stage_progress = data.get('stage_progress', {})
        journey_mode = data.get('journey_mode', 'guided')
        custom_workflow = data.get('custom_workflow')
        
        # حفظ التقدم في ملف مؤقت (نظراً لعدم وجود جدول journey_progress)
        progress_data = {
            'user_id': user_id,
            'template_id': template_id,
            'current_stage': current_stage,
            'completed_stages': completed_stages,
            'stage_progress': stage_progress,
            'journey_mode': journey_mode,
            'custom_workflow': custom_workflow,
            'updated_at': datetime.now().isoformat()
        }
        
        # حفظ في ملف مؤقت
        import os
        progress_dir = 'temp_files/journey_progress'
        os.makedirs(progress_dir, exist_ok=True)
        
        progress_file = f"{progress_dir}/{user_id}_{template_id}.json"
        with open(progress_file, 'w', encoding='utf-8') as f:
            json.dump(progress_data, f, ensure_ascii=False, indent=2)
        
        return jsonify({
            'success': True,
            'message': 'تم حفظ التقدم بنجاح'
        })
        
    except Exception as e:
        return jsonify({'error': f'خطأ في حفظ التقدم: {str(e)}'}), 500

def create_journey_workflow(template_id, journey_mode, start_choice, project_id):
    """إنشاء سير عمل للرحلة الموحدة"""
    
    # تحديد المراحل حسب اختيار البداية
    stage_mapping = {
        'analyze': [1, 2, 3, 4, 5, 6],
        'new': [2, 3, 4, 5, 6],
        'continue': [1, 2, 3, 4, 5, 6],  # سيتم تحديدها حسب التقدم
        'template': [1, 2, 3, 4, 5, 6]
    }
    
    selected_stages = stage_mapping.get(start_choice, [1, 2, 3, 4, 5, 6])
    
    # تعريف المراحل
    stage_definitions = {
        1: {
            'id': 1,
            'type': 'analyze_novel',
            'name': 'تحليل الرواية المصدر',
            'description': 'تحليل عميق للأسلوب والبنية والشخصيات',
            'required': start_choice == 'analyze',
            'customizable': journey_mode in ['custom', 'hybrid'],
            'auto_runnable': journey_mode == 'guided'
        },
        2: {
            'id': 2,
            'type': 'generate_ideas',
            'name': 'معمل الأفكار المحسن',
            'description': 'توليد أفكار مبدعة مبنية على التحليل',
            'required': True,
            'customizable': journey_mode in ['custom', 'hybrid'],
            'auto_runnable': journey_mode == 'guided'
        },
        3: {
            'id': 3,
            'type': 'build_blueprint',
            'name': 'باني المخطط الذكي',
            'description': 'بناء مخطط شامل للرواية الجديدة',
            'required': True,
            'customizable': journey_mode in ['custom', 'hybrid'],
            'auto_runnable': journey_mode == 'guided'
        },
        4: {
            'id': 4,
            'type': 'generate_chapter',
            'name': 'مولد الفصول الموجه',
            'description': 'توليد الفصول بالذكاء الاصطناعي',
            'required': True,
            'customizable': journey_mode in ['custom', 'hybrid'],
            'auto_runnable': journey_mode == 'guided'
        },
        5: {
            'id': 5,
            'type': 'refine_text',
            'name': 'المحرر التفاعلي المتقدم',
            'description': 'تحرير وتحسين الفصول بأدوات ذكية',
            'required': False,
            'customizable': True,
            'auto_runnable': False
        },
        6: {
            'id': 6,
            'type': 'generate_report',
            'name': 'التنقيح والتصدير النهائي',
            'description': 'مراجعة نهائية وتصدير المشروع',
            'required': True,
            'customizable': False,
            'auto_runnable': journey_mode == 'guided'
        }
    }
    
    # بناء سير العمل
    workflow_stages = []
    for i, stage_id in enumerate(selected_stages):
        stage = stage_definitions[stage_id].copy()
        stage.update({
            'position': {'x': 100, 'y': i * 120 + 100},
            'connections': [selected_stages[i + 1]] if i < len(selected_stages) - 1 else []
        })
        workflow_stages.append(stage)
    
    workflow = {
        'id': f'journey_{template_id}_{int(time.time())}',
        'name': f'رحلة موحدة - {journey_mode}',
        'description': f'رحلة كتابة موحدة بنمط {journey_mode}',
        'project_id': project_id,
        'mode': journey_mode,
        'start_choice': start_choice,
        'stages': workflow_stages,
        'metadata': {
            'created': datetime.now().isoformat(),
            'template_id': template_id,
            'estimated_total_time': calculate_workflow_time(workflow_stages),
            'total_stages': len(workflow_stages),
            'required_stages': len([s for s in workflow_stages if s.get('required', False)])
        }
    }
    
    return workflow

def calculate_workflow_time(stages):
    """حساب الوقت الإجمالي المتوقع لسير العمل"""
    time_estimates = {
        'analyze_novel': 15,
        'generate_ideas': 10,
        'build_blueprint': 20,
        'generate_chapter': 30,
        'refine_text': 45,
        'generate_report': 15
    }
    
    total_minutes = sum(time_estimates.get(stage.get('type', ''), 10) for stage in stages)
    
    hours = total_minutes // 60
    minutes = total_minutes % 60
    
    if hours > 0:
        return f"{hours} ساعة و {minutes} دقيقة"
    return f"{minutes} دقيقة"

# ================== APIs لوحة التحكم الشخصية ==================

@app.route('/api/style/profile', methods=['GET'])
def get_user_style_profile():
    """جلب ملف التفضيلات الشخصية للمستخدم"""
    try:
        adaptive = get_adaptive_service()
        user_id = adaptive.get_user_id(request)
        
        # جلب الملف الشخصي من قاعدة البيانات
        profile = adaptive.initialize_user_profile(user_id)
        
        if not profile:
            # إنشاء ملف افتراضي جديد
            profile = {
                'user_id': user_id,
                'jattlaoui_adaptation_level': 0.5,
                'preferred_vocabulary_complexity': 0.7,
                'preferred_sentence_length': 0.6,
                'preferred_cultural_depth': 0.8,
                'style_preferences': {
                    'dialogue_style': 'balanced',
                    'description_detail': 0.7,
                    'narrative_pace': 0.6,
                    'emotional_depth': 0.8,
                    'philosophical_focus': 0.6,
                    'sensory_description_intensity': 0.7
                },
                'writing_habits': {
                    'preferred_chapter_length': 3000,
                    'writing_sessions_per_week': 3,
                    'favorite_writing_time': 'evening'
                },
                'learning_stats': {
                    'total_interactions': 0,
                    'improvement_rate': 0.0,
                    'content_satisfaction': 3.5,
                    'last_activity': datetime.now().isoformat()
                }
            }
        
        return jsonify({
            'success': True,
            'profile': profile
        })
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': f"خطأ في جلب الملف الشخصي: {str(e)}"
        }), 500

@app.route('/api/style/profile', methods=['POST'])
def save_user_style_profile():
    """حفظ تفضيلات المستخدم الشخصية"""
    try:
        data = request.get_json()
        adaptive = get_adaptive_service()
        user_id = adaptive.get_user_id(request)
        
        # استخراج البيانات المطلوبة
        profile_updates = {}
        
        if 'jattlaoui_adaptation_level' in data:
            profile_updates['jattlaoui_adaptation_level'] = data['jattlaoui_adaptation_level']
        
        if 'preferred_vocabulary_complexity' in data:
            profile_updates['preferred_vocabulary_complexity'] = data['preferred_vocabulary_complexity']
        
        if 'preferred_sentence_length' in data:
            profile_updates['preferred_sentence_length'] = data['preferred_sentence_length']
        
        if 'preferred_cultural_depth' in data:
            profile_updates['preferred_cultural_depth'] = data['preferred_cultural_depth']
        
        if 'style_preferences' in data:
            profile_updates['style_preferences_json'] = json.dumps(data['style_preferences'], ensure_ascii=False)
        
        if 'writing_habits' in data:
            profile_updates['writing_habits_json'] = json.dumps(data['writing_habits'], ensure_ascii=False)
        
        # تحديث الملف الشخصي في قاعدة البيانات
        from database import update_writer_profile
        success = update_writer_profile(user_id, profile_updates)
        
        if success:
            return jsonify({
                'success': True,
                'message': 'تم حفظ التفضيلات بنجاح'
            })
        else:
            return jsonify({
                'success': False,
                'error': 'فشل في حفظ التفضيلات'
            }), 500
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': f"خطأ في حفظ التفضيلات: {str(e)}"
        }), 500

@app.route('/api/style/insights', methods=['GET'])
def get_learning_insights():
    """جلب رؤى التعلم ونشاط المستخدم"""
    try:
        adaptive = get_adaptive_service()
        user_id = adaptive.get_user_id(request)
        
        # جلب تحليل التفضيلات من قاعدة البيانات
        from database import analyze_user_preferences
        preferences_analysis = analyze_user_preferences(user_id)
        
        # تحليل البيانات وإنشاء رؤى
        insights = generate_user_insights(preferences_analysis, user_id)
        
        return jsonify({
            'success': True,
            'insights': insights
        })
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': f"خطأ في جلب رؤى التعلم: {str(e)}"
        }), 500

@app.route('/api/style/customized-prompt', methods=['POST'])
def get_customized_prompt():
    """إنشاء prompt مخصص بناءً على تفضيلات المستخدم"""
    try:
        data = request.get_json()
        content_type = data.get('content_type', 'general')
        base_prompt = data.get('base_prompt', '')
        
        adaptive = get_adaptive_service()
        user_id = adaptive.get_user_id(request)
        
        # الحصول على أسلوب الجطلاوي المخصص
        customized_style = adaptive.get_customized_jattlaoui_style(user_id)
        
        # توليد تعليمات مخصصة
        adaptive_instructions = adaptive.generate_adaptive_instructions(user_id, content_type)
        
        # دمج التعليمات مع prompt الأساسي
        enhanced_prompt = f"{base_prompt}\n\n{adaptive_instructions}"
        
        return jsonify({
            'success': True,
            'enhanced_prompt': enhanced_prompt,
            'customized_style': customized_style,
            'adaptive_instructions': adaptive_instructions
        })
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': f"خطأ في إنشاء prompt مخصص: {str(e)}"
        }), 500

@app.route('/api/style/rate-content', methods=['POST'])
def rate_generated_content():
    """تقييم المحتوى المولد لأغراض التعلم"""
    try:
        data = request.get_json()
        content_type = data.get('content_type', 'general')
        content = data.get('content', '')
        rating = data.get('rating', 3)  # من 1 إلى 5
        specific_feedback = data.get('specific_feedback', {})
        project_id = data.get('project_id')
        
        adaptive = get_adaptive_service()
        user_id = adaptive.get_user_id(request)
        
        # حفظ التقييم
        success = adaptive.save_user_rating(
            user_id, content_type, content, rating, specific_feedback, project_id
        )
        
        if success:
            # تطبيق التعلم من التفاعلات الجديدة
            adaptive.learn_from_interactions(user_id)
            
            return jsonify({
                'success': True,
                'message': 'تم حفظ التقييم وتحديث التفضيلات'
            })
        else:
            return jsonify({
                'success': False,
                'error': 'فشل في حفظ التقييم'
            }), 500
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': f"خطأ في حفظ التقييم: {str(e)}"
        }), 500

@app.route('/api/style/log-interaction', methods=['POST'])
def log_user_interaction():
    """تسجيل تفاعل المستخدم مع المحتوى"""
    try:
        data = request.get_json()
        interaction_type = data.get('interaction_type', 'content_edit')
        content_type = data.get('content_type', 'general')
        original_content = data.get('original_content', '')
        modified_content = data.get('modified_content', '')
        edit_details = data.get('edit_details', {})
        project_id = data.get('project_id')
        session_id = data.get('session_id')
        
        adaptive = get_adaptive_service()
        user_id = adaptive.get_user_id(request)
        
        # تسجيل التفاعل
        if interaction_type == 'content_edit':
            success = adaptive.log_content_edit(
                user_id, content_type, original_content, modified_content,
                edit_details, project_id, session_id
            )
        elif interaction_type == 'content_generated':
            success = adaptive.log_content_generation(
                user_id, content_type, modified_content, edit_details,
                project_id, session_id
            )
        else:
            return jsonify({
                'success': False,
                'error': 'نوع تفاعل غير مدعوم'
            }), 400
        
        if success:
            return jsonify({
                'success': True,
                'message': 'تم تسجيل التفاعل بنجاح'
            })
        else:
            return jsonify({
                'success': False,
                'error': 'فشل في تسجيل التفاعل'
            }), 500
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': f"خطأ في تسجيل التفاعل: {str(e)}"
        }), 500

def generate_user_insights(preferences_analysis: Dict[str, Any], user_id: str) -> Dict[str, Any]:
    """توليد رؤى المستخدم من تحليل التفضيلات"""
    try:
        # بيانات وهمية للتطوير - في التطبيق الحقيقي ستأتي من قاعدة البيانات
        current_date = datetime.now()
        insights = {
            'weekly_stats': {
                'content_generated': len(preferences_analysis.get('modification_patterns', [])),
                'content_expanded': sum(1 for p in preferences_analysis.get('modification_patterns', []) 
                                     if p.get('edit_type') == 'expansion'),
                'content_refined': sum(1 for p in preferences_analysis.get('modification_patterns', []) 
                                     if p.get('edit_type') == 'refinement'),
                'avg_satisfaction': 3.8
            },
            'style_evolution': [
                {
                    'date': (current_date - timedelta(days=i)).strftime('%Y-%m-%d'),
                    'vocabulary_complexity': 0.6 + (i * 0.02),
                    'sentence_length': 0.5 + (i * 0.03),
                    'cultural_depth': 0.7 + (i * 0.01),
                    'satisfaction': 3.5 + (i * 0.1)
                }
                for i in range(7, 0, -1)
            ],
            'content_distribution': [
                {'type': 'فصول', 'count': 12, 'avgRating': 4.2},
                {'type': 'حوارات', 'count': 8, 'avgRating': 3.9},
                {'type': 'أوصاف', 'count': 15, 'avgRating': 4.1},
                {'type': 'أفكار', 'count': 6, 'avgRating': 3.7}
            ],
            'improvement_patterns': [
                {'area': 'تعقيد المفردات', 'current': 0.7, 'target': 0.8},
                {'area': 'طول الجمل', 'current': 0.6, 'target': 0.7},
                {'area': 'العمق الثقافي', 'current': 0.8, 'target': 0.9},
                {'area': 'الوصف الحسي', 'current': 0.7, 'target': 0.8},
                {'area': 'التركيز الفلسفي', 'current': 0.6, 'target': 0.7}
            ]
        }
        
        return insights
        
    except Exception as e:
        print(f"خطأ في توليد رؤى المستخدم: {e}")
        return {
            'weekly_stats': {
                'content_generated': 0,
                'content_expanded': 0,
                'content_refined': 0,
                'avg_satisfaction': 3.5
            },
            'style_evolution': [],
            'content_distribution': [],
            'improvement_patterns': []
        }

# ================== APIs المحرر التفاعلي الذكي ==================

@app.route('/api/smart-editor/analyze', methods=['POST'])
def analyze_text_content():
    """تحليل شامل للنص مع اكتشاف المشاكل والإحصائيات"""
    try:
        data = request.get_json()
        text = data.get('text', '')
        user_profile = data.get('user_profile', {})
        analysis_type = data.get('analysis_type', 'comprehensive')
        
        if not text.strip():
            return jsonify({
                'success': False,
                'error': 'النص فارغ'
            }), 400
        
        # تحليل النص
        analysis = perform_text_analysis(text, user_profile, analysis_type)
        
        return jsonify({
            'success': True,
            'analysis': analysis
        })
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': f"خطأ في تحليل النص: {str(e)}"
        }), 500

@app.route('/api/smart-editor/suggestions', methods=['POST'])
def get_contextual_suggestions():
    """الحصول على اقتراحات سياقية ذكية للنص المحدد"""
    try:
        data = request.get_json()
        selected_text = data.get('selected_text', '')
        full_text = data.get('full_text', '')
        selection_context = data.get('selection_context', {})
        user_profile = data.get('user_profile', {})
        
        if not selected_text.strip():
            return jsonify({
                'success': False,
                'error': 'لم يتم تحديد نص'
            }), 400
        
        # توليد اقتراحات ذكية
        suggestions = generate_smart_suggestions(
            selected_text, full_text, selection_context, user_profile
        )
        
        return jsonify({
            'success': True,
            'suggestions': suggestions
        })
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': f"خطأ في توليد الاقتراحات: {str(e)}"
        }), 500

@app.route('/api/smart-editor/apply-suggestion', methods=['POST'])
def apply_smart_suggestion():
    """تطبيق اقتراح ذكي على النص"""
    try:
        data = request.get_json()
        original_text = data.get('original_text', '')
        selection = data.get('selection', {})
        suggestion = data.get('suggestion', {})
        user_profile = data.get('user_profile', {})
        
        if not original_text or not selection or not suggestion:
            return jsonify({
                'success': False,
                'error': 'بيانات غير مكتملة'
            }), 400
        
        # تطبيق الاقتراح
        result = apply_suggestion_to_text(
            original_text, selection, suggestion, user_profile
        )
        
        return jsonify({
            'success': True,
            'modified_text': result['modified_text'],
            'modification': result['modification'],
            'applied_changes': result['changes']
        })
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': f"خطأ في تطبيق الاقتراح: {str(e)}"
        }), 500

@app.route('/api/smart-editor/highlight-issues', methods=['POST'])
def highlight_text_issues():
    """تحديد وإبراز مشاكل النص للتغذية المرئية"""
    try:
        data = request.get_json()
        text = data.get('text', '')
        issue_types = data.get('issue_types', ['all'])
        
        if not text.strip():
            return jsonify({
                'success': False,
                'error': 'النص فارغ'
            }), 400
        
        # تحديد المشاكل
        issues = detect_text_issues(text, issue_types)
        
        return jsonify({
            'success': True,
            'issues': issues,
            'highlighted_ranges': [
                {
                    'start': issue['start'],
                    'end': issue['end'],
                    'type': issue['type'],
                    'severity': issue['severity']
                }
                for issue in issues
            ]
        })
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': f"خطأ في تحديد المشاكل: {str(e)}"
        }), 500

def perform_text_analysis(text: str, user_profile: dict, analysis_type: str) -> dict:
    """تحليل شامل للنص"""
    try:
        # إحصائيات أساسية
        words = text.split()
        sentences = text.count('.') + text.count('!') + text.count('?')
        paragraphs = len([p for p in text.split('\n') if p.strip()])
        
        # تحليل المشاكل
        issues = detect_text_issues(text)
        
        # حساب النقاط
        word_count = len(words)
        readability_score = calculate_readability_score(text)
        sentiment_score = calculate_sentiment_score(text)
        style_score = calculate_style_score(text, user_profile)
        overall_score = (readability_score + sentiment_score + style_score) / 3
        
        # تحديد مستوى التعقيد
        complexity_level = determine_complexity_level(text, word_count)
        
        return {
            'issues': issues,
            'statistics': {
                'wordCount': word_count,
                'sentenceCount': max(sentences, 1),
                'paragraphCount': max(paragraphs, 1),
                'readabilityScore': readability_score,
                'complexityLevel': complexity_level,
                'sentimentScore': sentiment_score,
                'styleScore': style_score
            },
            'suggestions': generate_improvement_suggestions(text, issues),
            'overallScore': overall_score
        }
        
    except Exception as e:
        print(f"خطأ في تحليل النص: {e}")
        return {
            'issues': [],
            'statistics': {
                'wordCount': 0,
                'sentenceCount': 0,
                'paragraphCount': 0,
                'readabilityScore': 0.5,
                'complexityLevel': 'متوسط',
                'sentimentScore': 0.5,
                'styleScore': 0.5
            },
            'suggestions': [],
            'overallScore': 0.5
        }

def detect_text_issues(text: str, issue_types: list = None) -> list:
    """اكتشاف مشاكل النص"""
    issues = []
    words = text.split()
    
    try:
        # اكتشاف التكرار
        word_counts = {}
        for i, word in enumerate(words):
            word_clean = word.strip('.,!?؛:').lower()
            if len(word_clean) > 3:  # تجاهل الكلمات القصيرة
                if word_clean in word_counts:
                    word_counts[word_clean].append(i)
                else:
                    word_counts[word_clean] = [i]
        
        # إضافة مشاكل التكرار
        for word, positions in word_counts.items():
            if len(positions) > 3:  # تكرار أكثر من 3 مرات
                for pos in positions[3:]:  # بدءاً من المرة الرابعة
                    start_pos = sum(len(words[j]) + 1 for j in range(pos))
                    end_pos = start_pos + len(words[pos])
                    
                    issues.append({
                        'id': f"rep_{pos}",
                        'type': 'repetition',
                        'severity': 'medium',
                        'start': start_pos,
                        'end': end_pos,
                        'text': words[pos],
                        'message': f"كلمة '{word}' مكررة كثيراً",
                        'suggestion': f"حاول استخدام مرادف لـ '{word}'"
                    })
        
        # اكتشاف الجمل الطويلة جداً
        sentences = text.split('.')
        current_pos = 0
        for i, sentence in enumerate(sentences):
            sentence = sentence.strip()
            if len(sentence.split()) > 25:  # جملة طويلة جداً
                issues.append({
                    'id': f"long_{i}",
                    'type': 'unclear',
                    'severity': 'medium',
                    'start': current_pos,
                    'end': current_pos + len(sentence),
                    'text': sentence[:50] + "...",
                    'message': "جملة طويلة قد تكون صعبة الفهم",
                    'suggestion': "قسم الجملة إلى جمل أقصر"
                })
            current_pos += len(sentence) + 1
        
        # اكتشاف علامات الترقيم المفقودة
        if not any(char in text for char in '.!?'):
            issues.append({
                'id': "punct_1",
                'type': 'grammar',
                'severity': 'high',
                'start': 0,
                'end': len(text),
                'text': text[:50] + "...",
                'message': "لا توجد علامات ترقيم في النص",
                'suggestion': "أضف نقاط وعلامات ترقيم مناسبة"
            })
        
        return issues[:10]  # إرجاع أول 10 مشاكل فقط
        
    except Exception as e:
        print(f"خطأ في اكتشاف مشاكل النص: {e}")
        return []

def generate_smart_suggestions(selected_text: str, full_text: str, 
                             selection_context: dict, user_profile: dict) -> list:
    """توليد اقتراحات ذكية للنص المحدد"""
    suggestions = []
    
    try:
        word_count = len(selected_text.split())
        
        # اقتراحات حسب طول النص
        if word_count < 10:
            suggestions.append({
                'id': 'expand_1',
                'type': 'expand',
                'title': 'توسيع النص',
                'description': 'إضافة تفاصيل أكثر لإثراء المعنى',
                'icon': 'Expand',
                'action': 'expand_text',
                'confidence': 0.8,
                'reasoning': 'النص قصير ويمكن إضافة تفاصيل'
            })
        elif word_count > 30:
            suggestions.append({
                'id': 'summarize_1',
                'type': 'summarize',
                'title': 'تلخيص النص',
                'description': 'اختصار النص مع الحفاظ على المعنى',
                'icon': 'Shrink',
                'action': 'summarize_text',
                'confidence': 0.7,
                'reasoning': 'النص طويل نسبياً ويمكن اختصاره'
            })
        
        # اقتراحات تحسين عامة
        suggestions.extend([
            {
                'id': 'improve_1',
                'type': 'improve',
                'title': 'تحسين الأسلوب',
                'description': 'تحسين اختيار الكلمات والتراكيب',
                'icon': 'Edit3',
                'action': 'improve_style',
                'confidence': 0.9,
                'reasoning': 'يمكن تحسين الأسلوب بشكل عام'
            },
            {
                'id': 'rephrase_1',
                'type': 'rephrase',
                'title': 'إعادة صياغة',
                'description': 'صياغة النص بطريقة مختلفة',
                'icon': 'RefreshCw',
                'action': 'rephrase_text',
                'confidence': 0.8,
                'reasoning': 'إعادة الصياغة قد تحسن الوضوح'
            }
        ])
        
        # اقتراحات مخصصة حسب ملف المستخدم
        if user_profile:
            vocab_level = user_profile.get('preferred_vocabulary_complexity', 0.5)
            if vocab_level > 0.7:
                suggestions.append({
                    'id': 'enhance_1',
                    'type': 'enhance',
                    'title': 'تعزيز المفردات',
                    'description': 'استخدام مفردات أكثر تقدماً',
                    'icon': 'Sparkles',
                    'action': 'enhance_vocabulary',
                    'confidence': 0.85,
                    'reasoning': 'يفضل المستخدم مفردات متقدمة'
                })
        
        return suggestions[:6]  # إرجاع أول 6 اقتراحات
        
    except Exception as e:
        print(f"خطأ في توليد الاقتراحات: {e}")
        return []

def apply_suggestion_to_text(original_text: str, selection: dict, 
                           suggestion: dict, user_profile: dict) -> dict:
    """تطبيق اقتراح على النص"""
    try:
        selected_text = selection.get('text', '')
        suggestion_type = suggestion.get('type', '')
        
        # محاكاة تطبيق الاقتراحات (في التطبيق الحقيقي ستكون هناك معالجة متقدمة)
        if suggestion_type == 'improve':
            modified_selection = improve_text_style(selected_text, user_profile)
        elif suggestion_type == 'expand':
            modified_selection = expand_text_content(selected_text, user_profile)
        elif suggestion_type == 'summarize':
            modified_selection = summarize_text_content(selected_text)
        elif suggestion_type == 'rephrase':
            modified_selection = rephrase_text_content(selected_text, user_profile)
        elif suggestion_type == 'enhance':
            modified_selection = enhance_vocabulary(selected_text, user_profile)
        else:
            modified_selection = selected_text
        
        # استبدال النص في النص الأصلي
        start_pos = selection.get('start', 0)
        end_pos = selection.get('end', len(selected_text))
        
        modified_text = (
            original_text[:original_text.find(selected_text)] + 
            modified_selection + 
            original_text[original_text.find(selected_text) + len(selected_text):]
        )
        
        return {
            'modified_text': modified_text,
            'modification': modified_selection,
            'changes': {
                'original': selected_text,
                'modified': modified_selection,
                'type': suggestion_type
            }
        }
        
    except Exception as e:
        print(f"خطأ في تطبيق الاقتراح: {e}")
        return {
            'modified_text': original_text,
            'modification': selection.get('text', ''),
            'changes': {}
        }

def improve_text_style(text: str, user_profile: dict) -> str:
    """تحسين أسلوب النص"""
    # محاكاة تحسين الأسلوب
    improvements = {
        'كان': 'بدا',
        'قال': 'صرح',
        'ذهب': 'توجه',
        'جميل': 'رائع',
        'كبير': 'عظيم'
    }
    
    improved_text = text
    for old, new in improvements.items():
        improved_text = improved_text.replace(old, new)
    
    return improved_text

def expand_text_content(text: str, user_profile: dict) -> str:
    """توسيع محتوى النص"""
    # محاكاة توسيع النص
    expansions = {
        'الليل': 'الليل البهيم',
        'النجوم': 'النجوم المتلألئة في السماء',
        'الذكريات': 'الذكريات العزيزة والغالية'
    }
    
    expanded_text = text
    for word, expansion in expansions.items():
        expanded_text = expanded_text.replace(word, expansion)
    
    return expanded_text

def summarize_text_content(text: str) -> str:
    """تلخيص محتوى النص"""
    # محاكاة تلخيص النص
    words = text.split()
    if len(words) > 10:
        return ' '.join(words[:len(words)//2]) + '...'
    return text

def rephrase_text_content(text: str, user_profile: dict) -> str:
    """إعادة صياغة النص"""
    # محاكاة إعادة الصياغة
    rephrasings = {
        'في ظلال': 'تحت ظلال',
        'يتأمل': 'يراقب',
        'المتلألئة': 'البراقة'
    }
    
    rephrased_text = text
    for old, new in rephrasings.items():
        rephrased_text = rephrased_text.replace(old, new)
    
    return rephrased_text

def enhance_vocabulary(text: str, user_profile: dict) -> str:
    """تعزيز المفردات"""
    # محاكاة تعزيز المفردات
    enhancements = {
        'جلس': 'استقر',
        'باردة': 'عليلة',
        'يداعب': 'يلامس'
    }
    
    enhanced_text = text
    for simple, advanced in enhancements.items():
        enhanced_text = enhanced_text.replace(simple, advanced)
    
    return enhanced_text

def calculate_readability_score(text: str) -> float:
    """حساب نقاط سهولة القراءة"""
    words = text.split()
    sentences = max(text.count('.') + text.count('!') + text.count('?'), 1)
    
    # صيغة مبسطة لحساب سهولة القراءة
    avg_words_per_sentence = len(words) / sentences
    avg_word_length = sum(len(word) for word in words) / len(words) if words else 0
    
    # كلما قلت الكلمات في الجملة وقل طول الكلمات، زادت سهولة القراءة
    readability = max(0, min(1, 1 - (avg_words_per_sentence / 30) - (avg_word_length / 15)))
    
    return readability

def calculate_sentiment_score(text: str) -> float:
    """حساب النقاط العاطفية"""
    # كلمات إيجابية وسلبية بسيطة
    positive_words = ['جميل', 'رائع', 'ممتاز', 'سعيد', 'فرح', 'حب']
    negative_words = ['سيء', 'حزين', 'ألم', 'صعب', 'مشكلة']
    
    words = text.lower().split()
    positive_count = sum(1 for word in words if any(pos in word for pos in positive_words))
    negative_count = sum(1 for word in words if any(neg in word for neg in negative_words))
    
    if positive_count + negative_count == 0:
        return 0.5  # محايد
    
    return positive_count / (positive_count + negative_count)

def calculate_style_score(text: str, user_profile: dict) -> float:
    """حساب نقاط الأسلوب"""
    # تقييم بسيط للأسلوب
    words = text.split()
    
    # تنوع المفردات
    unique_words = len(set(word.lower().strip('.,!?') for word in words))
    vocabulary_diversity = unique_words / len(words) if words else 0
    
    # طول الجمل
    sentences = max(text.count('.') + text.count('!') + text.count('?'), 1)
    avg_sentence_length = len(words) / sentences
    
    # نقاط الأسلوب (متوسط بسيط)
    style_score = (vocabulary_diversity + min(avg_sentence_length / 20, 1)) / 2
    
    return min(1, max(0, style_score))

def determine_complexity_level(text: str, word_count: int) -> str:
    """تحديد مستوى تعقيد النص"""
    sentences = max(text.count('.') + text.count('!') + text.count('?'), 1)
    avg_words_per_sentence = word_count / sentences
    
    if avg_words_per_sentence > 20:
        return 'متقدم'
    elif avg_words_per_sentence > 12:
        return 'متوسط'
    else:
        return 'بسيط'

def generate_improvement_suggestions(text: str, issues: list) -> list:
    """توليد اقتراحات التحسين العامة"""
    suggestions = []
    
    if len(issues) > 0:
        suggestions.append({
            'id': 'fix_issues',
            'type': 'correct',
            'title': 'إصلاح المشاكل',
            'description': f'إصلاح {len(issues)} مشكلة في النص',
            'icon': 'Check',
            'action': 'fix_all_issues',
            'confidence': 0.9,
            'reasoning': 'توجد مشاكل تحتاج إصلاح'
        })
    
    return suggestions

# ================== APIs الدمج السلس للأدوات والوكلاء ==================

@app.route('/api/tools/user-tools', methods=['GET'])
def get_user_tools():
    """الحصول على أدوات المستخدم المخصصة"""
    try:
        # في التطبيق الحقيقي، سيتم جلب الأدوات من قاعدة البيانات
        user_tools = [
            {
                'id': 'smart-editor',
                'name': 'المحرر الذكي الفائق',
                'type': 'utility',
                'description': 'محرر متقدم مع اقتراحات ذكية وتحليل فوري',
                'category': 'تحرير',
                'isActive': True,
                'isFavorite': True,
                'lastUsed': '2025-06-06T10:30:00Z',
                'version': '2.1.0',
                'status': 'running',
                'usage': {
                    'totalUses': 156,
                    'successRate': 96,
                    'avgExecutionTime': 1.2
                }
            },
            {
                'id': 'personal-style',
                'name': 'لوحة التحكم الشخصية',
                'type': 'template',
                'description': 'تخصيص أسلوب الكتابة والتفضيلات الشخصية',
                'category': 'تخصيص',
                'isActive': True,
                'isFavorite': True,
                'lastUsed': '2025-06-06T09:15:00Z',
                'version': '1.3.0',
                'status': 'running',
                'usage': {
                    'totalUses': 89,
                    'successRate': 94,
                    'avgExecutionTime': 0.8
                }
            },
            {
                'id': 'workflow-builder',
                'name': 'منشئ سير العمل',
                'type': 'utility',
                'description': 'بناء وإدارة سير عمل مخصص للكتابة',
                'category': 'إدارة',
                'isActive': True,
                'isFavorite': False,
                'lastUsed': '2025-06-05T16:45:00Z',
                'version': '1.2.1',
                'status': 'stopped',
                'usage': {
                    'totalUses': 34,
                    'successRate': 91,
                    'avgExecutionTime': 2.5
                }
            }
        ]
        
        return jsonify({
            'success': True,
            'tools': user_tools
        })
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': f"خطأ في جلب الأدوات: {str(e)}"
        }), 500

@app.route('/api/tools/quick-access', methods=['GET'])
def get_quick_access_tools():
    """الحصول على الأدوات للوصول السريع"""
    try:
        quick_tools = [
            {
                'id': 'smart-editor',
                'name': 'المحرر الذكي',
                'path': '/smart-editor',
                'category': 'تحرير',
                'isActive': True,
                'isFavorite': True,
                'hotkey': 'Ctrl+E'
            },
            {
                'id': 'personal-style',
                'name': 'التحكم الشخصي',
                'path': '/personal-style',
                'category': 'تخصيص',
                'isActive': True,
                'isFavorite': True,
                'hotkey': 'Ctrl+P'
            },
            {
                'id': 'agent-studio',
                'name': 'استوديو الوكلاء',
                'path': '/agent-studio',
                'category': 'وكلاء',
                'isActive': True,
                'isFavorite': False,
                'hotkey': 'Ctrl+A'
            },
            {
                'id': 'workflow-builder',
                'name': 'منشئ سير العمل',
                'path': '/workflow-builder',
                'category': 'إدارة',
                'isActive': True,
                'isFavorite': False,
                'hotkey': 'Ctrl+W'
            }
        ]
        
        return jsonify({
            'success': True,
            'tools': quick_tools
        })
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': f"خطأ في جلب أدوات الوصول السريع: {str(e)}"
        }), 500

@app.route('/api/tools/<tool_id>/favorite', methods=['POST'])
def toggle_tool_favorite(tool_id):
    """تبديل حالة المفضلة للأداة"""
    try:
        # في التطبيق الحقيقي، سيتم تحديث قاعدة البيانات
        return jsonify({
            'success': True,
            'message': 'تم تحديث المفضلة'
        })
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': f"خطأ في تحديث المفضلة: {str(e)}"
        }), 500

@app.route('/api/tools/<tool_id>/status', methods=['POST'])
def update_tool_status(tool_id):
    """تحديث حالة الأداة"""
    try:
        data = request.get_json()
        new_status = data.get('status', 'stopped')
        
        # في التطبيق الحقيقي، سيتم تحديث حالة الأداة
        return jsonify({
            'success': True,
            'message': f'تم تحديث حالة الأداة إلى {new_status}'
        })
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': f"خطأ في تحديث حالة الأداة: {str(e)}"
        }), 500

@app.route('/api/agents/user-agents', methods=['GET'])
def get_user_agents():
    """الحصول على وكلاء المستخدم"""
    try:
        user_agents = [
            {
                'id': 'idea-generator',
                'name': 'مولد الأفكار الإبداعية',
                'type': 'idea_generator_agent',
                'description': 'وكيل متخصص في توليد أفكار إبداعية ومبتكرة للكتابة',
                'state': 'idle',
                'capabilities': ['brainstorming', 'creative_thinking', 'idea_expansion'],
                'settings': {
                    'creativityLevel': 0.8,
                    'domainFocus': 'general'
                },
                'statistics': {
                    'tasksCompleted': 45,
                    'successRate': 92,
                    'averageResponseTime': 2.3
                }
            },
            {
                'id': 'cultural-maestro',
                'name': 'الخبير الثقافي',
                'type': 'cultural_maestro_agent',
                'description': 'وكيل متخصص في الجوانب الثقافية والتراثية العربية',
                'state': 'working',
                'currentTask': 'تحليل السياق الثقافي للنص',
                'capabilities': ['cultural_analysis', 'heritage_research', 'authenticity_check'],
                'settings': {
                    'culturalDepth': 0.9,
                    'historicalAccuracy': True
                },
                'statistics': {
                    'tasksCompleted': 78,
                    'successRate': 96,
                    'averageResponseTime': 1.8
                }
            },
            {
                'id': 'chapter-composer',
                'name': 'منسق الفصول',
                'type': 'chapter_composer_agent',
                'description': 'وكيل متخصص في تنسيق وهيكلة الفصول والمحتوى',
                'state': 'idle',
                'capabilities': ['chapter_structuring', 'content_organization', 'flow_optimization'],
                'settings': {
                    'structureComplexity': 0.7,
                    'flowOptimization': True
                },
                'statistics': {
                    'tasksCompleted': 23,
                    'successRate': 89,
                    'averageResponseTime': 3.5
                }
            }
        ]
        
        return jsonify({
            'success': True,
            'agents': user_agents
        })
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': f"خطأ في جلب الوكلاء: {str(e)}"
        }), 500

@app.route('/api/agents/active', methods=['GET'])
def get_active_agents():
    """الحصول على الوكلاء النشطين"""
    try:
        active_agents = [
            {
                'id': 'cultural-maestro',
                'name': 'الخبير الثقافي',
                'state': 'working',
                'type': 'cultural_maestro_agent',
                'currentTask': 'تحليل السياق الثقافي',
                'progress': 65
            },
            {
                'id': 'idea-generator',
                'name': 'مولد الأفكار',
                'state': 'idle',
                'type': 'idea_generator_agent'
            }
        ]
        
        return jsonify({
            'success': True,
            'agents': active_agents
        })
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': f"خطأ في جلب الوكلاء النشطين: {str(e)}"
        }), 500

@app.route('/api/agents/<agent_id>/status', methods=['GET'])
def get_agent_status(agent_id):
    """الحصول على حالة وكيل محدد"""
    try:
        # محاكاة حالة الوكيل
        agent_status = {
            'id': agent_id,
            'name': get_agent_name_by_id(agent_id),
            'state': 'working',
            'currentTask': 'تحليل النص وتوليد اقتراحات',
            'progress': 67,
            'performance': {
                'cpu': 15,
                'memory': 234,
                'tasksCompleted': 45,
                'errors': 2,
                'uptime': '2h 34m'
            },
            'settings': {
                'maxConcurrentTasks': 3,
                'autoRestart': True,
                'enableLogging': True,
                'priority': 'normal',
                'timeout': 30,
                'retryAttempts': 3,
                'creativityLevel': 0.7,
                'culturalDepth': 0.8,
                'responseSpeed': 0.6,
                'qualityThreshold': 0.8
            }
        }
        
        return jsonify({
            'success': True,
            'status': agent_status
        })
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': f"خطأ في جلب حالة الوكيل: {str(e)}"
        }), 500

@app.route('/api/agents/<agent_id>/control', methods=['POST'])
def control_agent(agent_id):
    """التحكم في الوكيل"""
    try:
        data = request.get_json()
        action = data.get('action')
        
        if action not in ['start', 'pause', 'stop', 'restart']:
            return jsonify({
                'success': False,
                'error': 'إجراء غير صالح'
            }), 400
        
        # في التطبيق الحقيقي، سيتم تنفيذ الإجراء على الوكيل الفعلي
        return jsonify({
            'success': True,
            'message': f'تم {action} الوكيل بنجاح',
            'new_state': get_new_agent_state(action)
        })
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': f"خطأ في التحكم بالوكيل: {str(e)}"
        }), 500

@app.route('/api/agents/<agent_id>/settings', methods=['PUT'])
def update_agent_settings(agent_id):
    """تحديث إعدادات الوكيل"""
    try:
        data = request.get_json()
        settings = data.get('settings', {})
        
        # في التطبيق الحقيقي، سيتم حفظ الإعدادات في قاعدة البيانات
        return jsonify({
            'success': True,
            'message': 'تم حفظ إعدادات الوكيل'
        })
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': f"خطأ في تحديث إعدادات الوكيل: {str(e)}"
        }), 500

@app.route('/api/agents/create-custom', methods=['POST'])
def create_custom_agent():
    """إنشاء وكيل مخصص"""
    try:
        data = request.get_json()
        
        # التحقق من صحة البيانات
        required_fields = ['name', 'description', 'capabilities']
        for field in required_fields:
            if field not in data or not data[field]:
                return jsonify({
                    'success': False,
                    'error': f'الحقل {field} مطلوب'
                }), 400
        
        # إنشاء معرف فريد للوكيل
        agent_id = f"custom_{data['name'].replace(' ', '_').lower()}"
        
        # في التطبيق الحقيقي، سيتم حفظ الوكيل في قاعدة البيانات وتهيئته
        new_agent = {
            'id': agent_id,
            'name': data['name'],
            'description': data['description'],
            'type': 'custom',
            'capabilities': data['capabilities'],
            'settings': data.get('settings', {}),
            'personalityTraits': data.get('personalityTraits', {}),
            'specialization': data.get('specialization', {}),
            'created_at': '2025-06-06T12:00:00Z',
            'status': 'created'
        }
        
        return jsonify({
            'success': True,
            'agent': new_agent,
            'message': f'تم إنشاء الوكيل "{data["name"]}" بنجاح'
        })
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': f"خطأ في إنشاء الوكيل: {str(e)}"
        }), 500

def get_agent_name_by_id(agent_id):
    """الحصول على اسم الوكيل بواسطة المعرف"""
    agent_names = {
        'idea-generator': 'مولد الأفكار الإبداعية',
        'cultural-maestro': 'الخبير الثقافي',
        'chapter-composer': 'منسق الفصول',
        'literary-critic': 'الناقد الأدبي',
        'novel-analyzer': 'محلل الروايات',
        'blueprint-architect': 'مهندس المخططات'
    }
    return agent_names.get(agent_id, 'وكيل ذكي')

def get_new_agent_state(action):
    """تحديد الحالة الجديدة للوكيل بناءً على الإجراء"""
    state_map = {
        'start': 'working',
        'pause': 'paused',
        'stop': 'stopped',
        'restart': 'working'
    }
    return state_map.get(action, 'idle')

# ================== تشغيل التطبيق ==================

# ==========================================
# APIs التحليلات الشخصية المتقدمة
# ==========================================

@app.route('/api/analytics/start-session', methods=['POST'])
def start_writing_session():
    """بدء جلسة كتابة جديدة"""
    try:
        if not AGENT_STUDIO_ENABLED:
            return jsonify({"error": "خدمة التحليلات غير متاحة"}), 503
            
        data = request.get_json()
        user_id = data.get('user_id', 'anonymous')
        project_id = data.get('project_id')
        stage_number = data.get('stage_number', 1)
        
        if not project_id:
            return jsonify({"error": "معرف المشروع مطلوب"}), 400
        
        session_id = analytics_service.start_writing_session(user_id, project_id, stage_number)
        
        return jsonify({
            "success": True,
            "session_id": session_id,
            "message": "تم بدء جلسة الكتابة بنجاح"
        })
        
    except Exception as e:
        print(f"خطأ في بدء جلسة الكتابة: {str(e)}")
        return jsonify({"error": f"خطأ في بدء الجلسة: {str(e)}"}), 500

@app.route('/api/analytics/end-session', methods=['POST'])
def end_writing_session():
    """إنهاء جلسة كتابة وحفظ الإحصائيات"""
    try:
        if not AGENT_STUDIO_ENABLED:
            return jsonify({"error": "خدمة التحليلات غير متاحة"}), 503
            
        data = request.get_json()
        user_id = data.get('user_id', 'anonymous')
        session_id = data.get('session_id')
        session_data = data.get('session_data', {})
        
        if not session_id:
            return jsonify({"error": "معرف الجلسة مطلوب"}), 400
        
        analytics_service.end_writing_session(user_id, session_id, session_data)
        
        return jsonify({
            "success": True,
            "message": "تم إنهاء جلسة الكتابة وحفظ الإحصائيات"
        })
        
    except Exception as e:
        print(f"خطأ في إنهاء جلسة الكتابة: {str(e)}")
        return jsonify({"error": f"خطأ في إنهاء الجلسة: {str(e)}"}), 500

@app.route('/api/analytics/analyze-text', methods=['POST'])
def analyze_text_style():
    """تحليل أسلوب النص وحفظ النتائج"""
    try:
        if not AGENT_STUDIO_ENABLED:
            return jsonify({"error": "خدمة التحليلات غير متاحة"}), 503
            
        data = request.get_json()
        user_id = data.get('user_id', 'anonymous')
        content = data.get('content', '')
        content_type = data.get('content_type', 'chapter')
        project_id = data.get('project_id')
        
        if not content.strip():
            return jsonify({"error": "المحتوى مطلوب للتحليل"}), 400
        
        analysis = analytics_service.analyze_text_style(user_id, content, content_type, project_id)
        
        return jsonify({
            "success": True,
            "analysis": analysis,
            "message": "تم تحليل النص بنجاح"
        })
        
    except Exception as e:
        print(f"خطأ في تحليل النص: {str(e)}")
        return jsonify({"error": f"خطأ في تحليل النص: {str(e)}"}), 500

@app.route('/api/analytics/progress', methods=['GET'])
def get_progress_analytics():
    """جلب تحليلات تقدم الكتابة"""
    try:
        if not AGENT_STUDIO_ENABLED:
            return jsonify({"error": "خدمة التحليلات غير متاحة"}), 503
            
        user_id = request.args.get('user_id', 'anonymous')
        project_id = request.args.get('project_id', type=int)
        days = request.args.get('days', 30, type=int)
        
        analytics = analytics_service.get_writing_progress_analytics(user_id, project_id, days)
        
        return jsonify({
            "success": True,
            "analytics": analytics,
            "message": "تم جلب تحليلات التقدم بنجاح"
        })
        
    except Exception as e:
        print(f"خطأ في جلب تحليلات التقدم: {str(e)}")
        return jsonify({"error": f"خطأ في جلب التحليلات: {str(e)}"}), 500

@app.route('/api/analytics/personal-report/<int:project_id>', methods=['GET'])
def get_personalized_report(project_id):
    """إنتاج تقرير شخصي مفصل عن الرواية"""
    try:
        if not AGENT_STUDIO_ENABLED:
            return jsonify({"error": "خدمة التحليلات غير متاحة"}), 503
            
        user_id = request.args.get('user_id', 'anonymous')
        
        report = analytics_service.generate_personalized_report(user_id, project_id)
        
        if 'error' in report:
            return jsonify(report), 404
        
        return jsonify({
            "success": True,
            "report": report,
            "message": "تم إنتاج التقرير الشخصي بنجاح"
        })
        
    except Exception as e:
        print(f"خطأ في إنتاج التقرير الشخصي: {str(e)}")
        return jsonify({"error": f"خطأ في إنتاج التقرير: {str(e)}"}), 500

@app.route('/api/analytics/style-evolution', methods=['GET'])
def get_style_evolution():
    """جلب تطور الأسلوب الشخصي عبر الوقت"""
    try:
        if not AGENT_STUDIO_ENABLED:
            return jsonify({"error": "خدمة التحليلات غير متاحة"}), 503
            
        user_id = request.args.get('user_id', 'anonymous')
        project_id = request.args.get('project_id', type=int)
        
        conn = get_db_connection()
        cursor = conn.cursor()
        
        query = '''
            SELECT analysis_date, metaphor_density, vocabulary_complexity, 
                   formality_score, creativity_score, coherence_score,
                   avg_sentence_length, cultural_references_count
            FROM style_analysis 
            WHERE user_identifier = ?
        '''
        params = [user_id]
        
        if project_id:
            query += ' AND project_id = ?'
            params.append(project_id)
        
        query += ' ORDER BY analysis_date'
        
        cursor.execute(query, params)
        evolution_data = cursor.fetchall()
        conn.close()
        
        # تحويل البيانات إلى تنسيق مناسب للرسوم البيانية
        evolution = []
        for row in evolution_data:
            evolution.append({
                'date': row['analysis_date'],
                'metaphor_density': row['metaphor_density'],
                'vocabulary_complexity': row['vocabulary_complexity'],
                'formality_score': row['formality_score'],
                'creativity_score': row['creativity_score'],
                'coherence_score': row['coherence_score'],
                'avg_sentence_length': row['avg_sentence_length'],
                'cultural_references': row['cultural_references_count']
            })
        
        return jsonify({
            "success": True,
            "evolution": evolution,
            "message": "تم جلب تطور الأسلوب بنجاح"
        })
        
    except Exception as e:
        print(f"خطأ في جلب تطور الأسلوب: {str(e)}")
        return jsonify({"error": f"خطأ في جلب البيانات: {str(e)}"}), 500

@app.route('/api/analytics/writing-sessions', methods=['GET'])
def get_writing_sessions():
    """جلب إحصائيات جلسات الكتابة"""
    try:
        if not AGENT_STUDIO_ENABLED:
            return jsonify({"error": "خدمة التحليلات غير متاحة"}), 503
            
        user_id = request.args.get('user_id', 'anonymous')
        project_id = request.args.get('project_id', type=int)
        days = request.args.get('days', 30, type=int)
        
        conn = get_db_connection()
        cursor = conn.cursor()
        
        # فترة التحليل
        from datetime import datetime, timedelta
        start_date = (datetime.now() - timedelta(days=days)).isoformat()
        
        query = '''
            SELECT session_start, session_duration, words_written, 
                   edits_count, quality_score, stage_number
            FROM writing_sessions 
            WHERE user_identifier = ? AND session_start >= ?
        '''
        params = [user_id, start_date]
        
        if project_id:
            query += ' AND project_id = ?'
            params.append(project_id)
        
        query += ' ORDER BY session_start'
        
        cursor.execute(query, params)
        sessions_data = cursor.fetchall()
        conn.close()
        
        # تحويل البيانات
        sessions = []
        for row in sessions_data:
            sessions.append({
                'date': row['session_start'],
                'duration': row['session_duration'],
                'words': row['words_written'],
                'edits': row['edits_count'],
                'quality': row['quality_score'],
                'stage': row['stage_number']
            })
        
        return jsonify({
            "success": True,
            "sessions": sessions,
            "message": "تم جلب بيانات الجلسات بنجاح"
        })
        
    except Exception as e:
        print(f"خطأ في جلب جلسات الكتابة: {str(e)}")
        return jsonify({"error": f"خطأ في جلب البيانات: {str(e)}"}), 500

@app.route('/api/analytics/dashboard-stats', methods=['GET'])
def get_dashboard_stats():
    """جلب إحصائيات لوحة التحكم الرئيسية"""
    try:
        if not AGENT_STUDIO_ENABLED:
            return jsonify({"error": "خدمة التحليلات غير متاحة"}), 503
            
        user_id = request.args.get('user_id', 'anonymous')
        
        # جلب الإحصائيات الأساسية
        analytics_30_days = analytics_service.get_writing_progress_analytics(user_id, None, 30)
        analytics_7_days = analytics_service.get_writing_progress_analytics(user_id, None, 7)
        
        # إحصائيات مقارنة
        stats = {
            'current_week': {
                'sessions': analytics_7_days.get('total_sessions', 0),
                'words': analytics_7_days.get('total_words', 0),
                'quality': analytics_7_days.get('average_quality', 0),
                'consistency': analytics_7_days.get('writing_consistency', 0)
            },
            'current_month': {
                'sessions': analytics_30_days.get('total_sessions', 0),
                'words': analytics_30_days.get('total_words', 0),
                'quality': analytics_30_days.get('average_quality', 0),
                'consistency': analytics_30_days.get('writing_consistency', 0)
            },
            'productivity': analytics_30_days.get('productivity_stats', {}),
            'improvement_areas': analytics_30_days.get('improvement_areas', []),
            'style_evolution': analytics_30_days.get('style_evolution', {})
        }
        
        return jsonify({
            "success": True,
            "stats": stats,
            "message": "تم جلب إحصائيات لوحة التحكم بنجاح"
        })
        
    except Exception as e:
        print(f"خطأ في جلب إحصائيات لوحة التحكم: {str(e)}")
        return jsonify({"error": f"خطأ في جلب الإحصائيات: {str(e)}"}), 500

# ===== APIs معالجة PDF المتقدمة =====

@app.route('/api/pdf/info', methods=['POST'])
def get_pdf_info():
    """الحصول على معلومات مفصلة عن ملف PDF"""
    try:
        if not PDF_SERVICE_AVAILABLE:
            return jsonify({
                "error": "خدمة معالجة PDF المتقدمة غير متاحة",
                "fallback_available": True
            }), 503
        
        if 'file' not in request.files:
            return jsonify({"error": "لم يتم رفع ملف"}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({"error": "لم يتم اختيار ملف"}), 400
        
        # التحقق من نوع الملف
        if not (file.mimetype == 'application/pdf' or file.filename.endswith('.pdf')):
            return jsonify({"error": "يجب أن يكون الملف من نوع PDF"}), 400
        
        # قراءة بيانات الملف
        pdf_data = file.read()
        
        # الحصول على معلومات PDF
        pdf_service = get_pdf_service()
        pdf_info = pdf_service.get_pdf_info(pdf_data)
        
        return jsonify({
            "success": True,
            "info": pdf_info,
            "message": "تم تحليل ملف PDF بنجاح"
        })
        
    except Exception as e:
        print(f"خطأ في تحليل معلومات PDF: {str(e)}")
        return jsonify({"error": f"خطأ في معالجة الملف: {str(e)}"}), 500

@app.route('/api/pdf/extract-advanced', methods=['POST'])
def extract_pdf_advanced():
    """استخراج متقدم لمحتوى PDF مع جداول وصور"""
    try:
        if not PDF_SERVICE_AVAILABLE:
            return jsonify({
                "error": "خدمة معالجة PDF المتقدمة غير متاحة",
                "fallback_available": True
            }), 503
        
        if 'file' not in request.files:
            return jsonify({"error": "لم يتم رفع ملف"}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({"error": "لم يتم اختيار ملف"}), 400
        
        # التحقق من نوع الملف
        if not (file.mimetype == 'application/pdf' or file.filename.endswith('.pdf')):
            return jsonify({"error": "يجب أن يكون الملف من نوع PDF"}), 400
        
        # قراءة المعاملات
        extract_tables = request.form.get('extract_tables', 'true').lower() == 'true'
        extract_images = request.form.get('extract_images', 'false').lower() == 'true'
        extraction_method = request.form.get('method', 'auto')
        
        # قراءة بيانات الملف
        pdf_data = file.read()
        
        # الاستخراج المتقدم
        pdf_service = get_pdf_service()
        
        try:
            from advanced_pdf_service import ExtractionMethod
            method = ExtractionMethod(extraction_method) if extraction_method != 'auto' else ExtractionMethod.AUTO
        except ValueError:
            method = ExtractionMethod.AUTO
        
        result = pdf_service.extract_pdf_content(
            pdf_data, 
            method=method,
            extract_tables=extract_tables,
            extract_images=extract_images
        )
        
        # تحويل النتيجة لتنسيق JSON
        response_data = {
            "success": result.success,
            "text": result.text,
            "metadata": {
                "title": result.metadata.title,
                "author": result.metadata.author,
                "subject": result.metadata.subject,
                "pages_count": result.metadata.pages_count,
                "has_arabic_text": result.metadata.has_arabic_text,
                "file_size": result.metadata.file_size,
                "is_encrypted": result.metadata.is_encrypted
            },
            "pages_count": len(result.page_texts),
            "tables_count": len(result.tables),
            "images_count": len(result.images),
            "extraction_method": result.extraction_method,
            "processing_time": result.processing_time,
            "error_message": result.error_message
        }
        
        # إضافة الجداول إذا طُلبت
        if extract_tables and result.tables:
            response_data["tables"] = [
                {
                    "page_number": table.page_number,
                    "table_index": table.table_index,
                    "rows_count": len(table.data),
                    "columns_count": len(table.data[0]) if table.data else 0,
                    "data": table.data[:10]  # أول 10 صفوف فقط للمعاينة
                }
                for table in result.tables[:5]  # أول 5 جداول فقط
            ]
        
        # إضافة معلومات الصور إذا طُلبت (بدون البيانات الفعلية)
        if extract_images and result.images:
            response_data["images"] = [
                {
                    "page_number": img.page_number,
                    "image_index": img.image_index,
                    "format": img.image_format,
                    "size": img.size,
                    "data_size": len(img.image_data)
                }
                for img in result.images[:10]  # أول 10 صور فقط
            ]
        
        return jsonify(response_data)
        
    except Exception as e:
        print(f"خطأ في الاستخراج المتقدم للPDF: {str(e)}")
        return jsonify({"error": f"خطأ في معالجة الملف: {str(e)}"}), 500

@app.route('/api/pdf/methods', methods=['GET'])
def get_pdf_methods():
    """الحصول على طرق استخراج PDF المتاحة"""
    try:
        if PDF_SERVICE_AVAILABLE:
            pdf_service = get_pdf_service()
            return jsonify({
                "success": True,
                "available_methods": {
                    method.value: available 
                    for method, available in pdf_service.available_methods.items()
                },
                "preferred_order": [method.value for method in pdf_service.preferred_order],
                "advanced_service": True
            })
        else:
            return jsonify({
                "success": True,
                "available_methods": {"pypdf2": True},
                "preferred_order": ["pypdf2"],
                "advanced_service": False,
                "message": "الخدمة المتقدمة غير متاحة، استخدام PyPDF2 فقط"
            })
            
    except Exception as e:
        return jsonify({"error": f"خطأ في فحص الطرق المتاحة: {str(e)}"}), 500

@app.route('/api/pdf/test', methods=['POST'])
def test_pdf_extraction():
    """اختبار طرق استخراج PDF المختلفة"""
    try:
        if not PDF_SERVICE_AVAILABLE:
            return jsonify({
                "error": "خدمة معالجة PDF المتقدمة غير متاحة"
            }), 503
        
        if 'file' not in request.files:
            return jsonify({"error": "لم يتم رفع ملف"}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({"error": "لم يتم اختيار ملف"}), 400
        
        # التحقق من نوع الملف
        if not (file.mimetype == 'application/pdf' or file.filename.endswith('.pdf')):
            return jsonify({"error": "يجب أن يكون الملف من نوع PDF"}), 400
        
        # قراءة بيانات الملف
        pdf_data = file.read()
        
        pdf_service = get_pdf_service()
        results = {}
        
        # تجربة كل طريقة متاحة
        from advanced_pdf_service import ExtractionMethod
        methods_to_test = [
            ExtractionMethod.PYMUPDF,
            ExtractionMethod.PDFPLUMBER,
            ExtractionMethod.PDFMINER,
            ExtractionMethod.PYPDF2
        ]
        
        for method in methods_to_test:
            if pdf_service.available_methods.get(method, False):
                try:
                    result = pdf_service.extract_pdf_content(pdf_data, method=method, extract_tables=False, extract_images=False)
                    results[method.value] = {
                        "success": result.success,
                        "text_length": len(result.text),
                        "pages_count": len(result.page_texts),
                        "processing_time": result.processing_time,
                        "has_arabic_text": result.metadata.has_arabic_text,
                        "error_message": result.error_message
                    }
                except Exception as e:
                    results[method.value] = {
                        "success": False,
                        "error_message": str(e)
                    }
        
        return jsonify({
            "success": True,
            "test_results": results,
            "file_info": {
                "filename": file.filename,
                "size": len(pdf_data)
            }
        })
        
    except Exception as e:
        print(f"خطأ في اختبار استخراج PDF: {str(e)}")
        return jsonify({"error": f"خطأ في الاختبار: {str(e)}"}), 500

# ==============================================================================
# APIs المحرر التفاعلي الذكي الفائق - Advanced Text Processing APIs
# ==============================================================================

try:
    from advanced_text_processing_service import AdvancedTextProcessingService
    text_processor = AdvancedTextProcessingService()
    TEXT_PROCESSING_ENABLED = True
    print("✅ خدمة معالجة النصوص المتقدمة متاحة")
except ImportError as e:
    TEXT_PROCESSING_ENABLED = False
    print(f"⚠️ خدمة معالجة النصوص المتقدمة غير متاحة: {e}")

@app.route('/api/text/analyze', methods=['POST'])
def analyze_text_comprehensive():
    """تحليل شامل للنص - جودة، إحصائيات، أسلوب، تدفق"""
    if not TEXT_PROCESSING_ENABLED:
        return jsonify({"error": "خدمة معالجة النصوص غير متاحة"}), 503
    
    try:
        data = request.json
        text = data.get('text', '')
        
        if not text or len(text.strip()) < 10:
            return jsonify({"error": "النص قصير جداً للتحليل"}), 400
        
        # تشغيل التحليل الشامل
        loop = asyncio.new_event_loop()
        asyncio.set_event_loop(loop)
        analysis_result = loop.run_until_complete(
            text_processor.analyze_text_comprehensive(text)
        )
        loop.close()
        
        return jsonify({
            "success": True,
            "analysis": analysis_result
        })
        
    except Exception as e:
        return jsonify({
            "error": f"خطأ في تحليل النص: {str(e)}"
        }), 500

@app.route('/api/text/adjust-length', methods=['POST'])
def adjust_text_length():
    """تعديل طول النص (إطالة أو تقصير) مع الحفاظ على المعنى"""
    if not TEXT_PROCESSING_ENABLED:
        return jsonify({"error": "خدمة معالجة النصوص غير متاحة"}), 503
    
    try:
        data = request.json
        text = data.get('text', '')
        target_length = data.get('target_length', 0)
        operation = data.get('operation', 'auto')  # 'shorten', 'expand', 'auto'
        
        if not text or len(text.strip()) < 10:
            return jsonify({"error": "النص قصير جداً للمعالجة"}), 400
        
        if target_length <= 0:
            return jsonify({"error": "الطول المطلوب يجب أن يكون أكبر من صفر"}), 400
        
        # تعديل طول النص
        loop = asyncio.new_event_loop()
        asyncio.set_event_loop(loop)
        adjustment_result = loop.run_until_complete(
            text_processor.adjust_text_length(text, target_length, operation)
        )
        loop.close()
        
        return jsonify({
            "success": True,
            "result": adjustment_result
        })
        
    except Exception as e:
        return jsonify({
            "error": f"خطأ في تعديل طول النص: {str(e)}"
        }), 500

@app.route('/api/text/suggestions', methods=['POST'])
def generate_text_suggestions():
    """توليد اقتراحات تحسين للنص"""
    if not TEXT_PROCESSING_ENABLED:
        return jsonify({"error": "خدمة معالجة النصوص غير متاحة"}), 503
    
    try:
        data = request.json
        text = data.get('text', '')
        suggestion_type = data.get('type', 'improve')  # 'improve', 'rephrase', 'enhance', 'all'
        
        if not text or len(text.strip()) < 10:
            return jsonify({"error": "النص قصير جداً لتوليد الاقتراحات"}), 400
        
        # توليد الاقتراحات
        loop = asyncio.new_event_loop()
        asyncio.set_event_loop(loop)
        suggestions = loop.run_until_complete(
            text_processor.generate_text_suggestions(text, suggestion_type)
        )
        loop.close()
        
        return jsonify({
            "success": True,
            "suggestions": [suggestion.__dict__ for suggestion in suggestions]
        })
        
    except Exception as e:
        return jsonify({
            "error": f"خطأ في توليد الاقتراحات: {str(e)}"
        }), 500

@app.route('/api/text/quality-check', methods=['POST'])
def check_text_quality():
    """فحص سريع لجودة النص مع توصيات"""
    if not TEXT_PROCESSING_ENABLED:
        return jsonify({"error": "خدمة معالجة النصوص غير متاحة"}), 503
    
    try:
        data = request.json
        text = data.get('text', '')
        
        if not text or len(text.strip()) < 10:
            return jsonify({"error": "النص قصير جداً للفحص"}), 400
        
        # تشغيل فحص الجودة
        loop = asyncio.new_event_loop()
        asyncio.set_event_loop(loop)
        quality_analysis = loop.run_until_complete(
            text_processor._analyze_text_quality(text)
        )
        loop.close()
        
        # فحص المشاكل الأساسية
        issues = text_processor._detect_text_issues(text)
        
        return jsonify({
            "success": True,
            "quality": quality_analysis,
            "issues": [issue.__dict__ for issue in issues],
            "word_count": len(text.split()),
            "character_count": len(text)
        })
        
    except Exception as e:
        return jsonify({
            "error": f"خطأ في فحص جودة النص: {str(e)}"
        }), 500

@app.route('/api/text/metrics', methods=['POST'])
def get_text_metrics():
    """الحصول على مقاييس النص الأساسية"""
    if not TEXT_PROCESSING_ENABLED:
        return jsonify({"error": "خدمة معالجة النصوص غير متاحة"}), 503
    
    try:
        data = request.json
        text = data.get('text', '')
        
        if not text:
            return jsonify({"error": "النص مفقود"}), 400
        
        # حساب المقاييس
        metrics = text_processor._calculate_text_metrics(text)
        
        return jsonify({
            "success": True,
            "metrics": metrics.__dict__
        })
        
    except Exception as e:
        return jsonify({
            "error": f"خطأ في حساب مقاييس النص: {str(e)}"
        }), 500

@app.route('/api/text/style-analysis', methods=['POST'])
def analyze_writing_style():
    """تحليل أسلوب الكتابة والنبرة"""
    if not TEXT_PROCESSING_ENABLED:
        return jsonify({"error": "خدمة معالجة النصوص غير متاحة"}), 503
    
    try:
        data = request.json
        text = data.get('text', '')
        
        if not text or len(text.strip()) < 50:
            return jsonify({"error": "النص قصير جداً لتحليل الأسلوب"}), 400
        
        # تحليل الأسلوب
        style_analysis = text_processor._analyze_writing_style(text)
        
        # تحليل التدفق
        flow_analysis = text_processor._analyze_text_flow(text)
        
        # تحليل المفردات
        vocabulary_analysis = text_processor._analyze_vocabulary(text)
        
        return jsonify({
            "success": True,
            "style": style_analysis,
            "flow": flow_analysis,
            "vocabulary": vocabulary_analysis
        })
        
    except Exception as e:
        return jsonify({
            "error": f"خطأ في تحليل أسلوب الكتابة: {str(e)}"
        }), 500

@app.route('/api/text/readability', methods=['POST'])
def check_text_readability():
    """فحص قابلية القراءة والفهم"""
    if not TEXT_PROCESSING_ENABLED:
        return jsonify({"error": "خدمة معالجة النصوص غير متاحة"}), 503
    
    try:
        data = request.json
        text = data.get('text', '')
        
        if not text or len(text.strip()) < 20:
            return jsonify({"error": "النص قصير جداً لفحص القابلية للقراءة"}), 400
        
        # حساب المقاييس الأساسية
        metrics = text_processor._calculate_text_metrics(text)
        
        # تحديد مستوى التعقيد
        if metrics.complexity_score <= 0.3:
            complexity_level = "بسيط"
            complexity_description = "سهل الفهم والقراءة"
        elif metrics.complexity_score <= 0.6:
            complexity_level = "متوسط"
            complexity_description = "يتطلب تركيز متوسط"
        elif metrics.complexity_score <= 0.8:
            complexity_level = "معقد"
            complexity_description = "يتطلب تركيز عالي"
        else:
            complexity_level = "متقدم"
            complexity_description = "يتطلب خبرة متقدمة"
        
        # تحديد مستوى القابلية للقراءة
        if metrics.readability_score >= 0.8:
            readability_level = "ممتاز"
        elif metrics.readability_score >= 0.6:
            readability_level = "جيد"
        elif metrics.readability_score >= 0.4:
            readability_level = "مقبول"
        else:
            readability_level = "يحتاج تحسين"
        
        return jsonify({
            "success": True,
            "readability": {
                "score": metrics.readability_score,
                "level": readability_level,
                "reading_time_minutes": metrics.reading_time_minutes,
                "complexity": {
                    "score": metrics.complexity_score,
                    "level": complexity_level,
                    "description": complexity_description
                },
                "metrics": {
                    "average_sentence_length": metrics.average_sentence_length,
                    "word_count": metrics.word_count,
                    "sentence_count": metrics.sentence_count
                }
            }
        })
        
    except Exception as e:
        return jsonify({
            "error": f"خطأ في فحص قابلية القراءة: {str(e)}"
        }), 500

# ==============================================================================
# APIs التحكم الذكي بطول النص - Intelligent Text Length Control APIs  
# ==============================================================================

@app.route('/api/text/modify-length', methods=['POST'])
def modify_text_length():
    """تعديل طول النص بذكاء - إطالة أو تقصير أو تحسين"""
    if not TEXT_PROCESSING_ENABLED:
        return jsonify({"error": "خدمة معالجة النصوص غير متاحة"}), 503
    
    try:
        data = request.json
        text = data.get('text', '')
        modification_type = data.get('modification_type', 'EXPAND')  # EXPAND, SUMMARIZE, IMPROVE, REPHRASE
        target_length = data.get('target_length')  # عدد الكلمات المطلوبة
        style_preferences = data.get('style_preferences', {})
        context = data.get('context', '')
        preserve_style = data.get('preserve_style', True)
        
        if not text or len(text.strip()) < 10:
            return jsonify({"error": "النص قصير جداً للمعالجة"}), 400
        
        # إنشاء طلب التعديل
        from advanced_text_processing_service import TextModificationRequest, TextModificationType
        
        # تحويل نوع التعديل
        modification_type_enum = {
            'EXPAND': TextModificationType.EXPAND,
            'SUMMARIZE': TextModificationType.SUMMARIZE,
            'IMPROVE': TextModificationType.IMPROVE,
            'REPHRASE': TextModificationType.REPHRASE,
            'SIMPLIFY': TextModificationType.SIMPLIFY,
            'ENHANCE': TextModificationType.ENHANCE
        }.get(modification_type, TextModificationType.IMPROVE)
        
        request_obj = TextModificationRequest(
            text=text,
            modification_type=modification_type_enum,
            target_length=target_length,
            style_preferences=style_preferences,
            context=context,
            preserve_style=preserve_style
        )
        
        # تشغيل التعديل
        loop = asyncio.new_event_loop()
        asyncio.set_event_loop(loop)
        modification_result = loop.run_until_complete(
            text_processor.modify_text_length(request_obj)
        )
        loop.close()
        
        # تحويل النتيجة لصيغة JSON
        result_data = {
            "success": True,
            "modification": {
                "original_text": modification_result.original_text,
                "modified_text": modification_result.modified_text,
                "modification_type": modification_result.modification_type.value,
                "quality_score": modification_result.quality_score,
                "original_metrics": {
                    "word_count": modification_result.original_metrics.word_count,
                    "sentence_count": modification_result.original_metrics.sentence_count,
                    "character_count": modification_result.original_metrics.character_count,
                    "reading_time_minutes": modification_result.original_metrics.reading_time_minutes,
                    "complexity_score": modification_result.original_metrics.complexity_score,
                    "readability_score": modification_result.original_metrics.readability_score
                },
                "new_metrics": {
                    "word_count": modification_result.new_metrics.word_count,
                    "sentence_count": modification_result.new_metrics.sentence_count,
                    "character_count": modification_result.new_metrics.character_count,
                    "reading_time_minutes": modification_result.new_metrics.reading_time_minutes,
                    "complexity_score": modification_result.new_metrics.complexity_score,
                    "readability_score": modification_result.new_metrics.readability_score
                },
                "changes_summary": modification_result.changes_summary,
                "suggestions": [
                    {
                        "id": suggestion.id,
                        "type": suggestion.type,
                        "original_text": suggestion.original_text,
                        "suggested_text": suggestion.suggested_text,
                        "reasoning": suggestion.reasoning,
                        "confidence": suggestion.confidence,
                        "impact_score": suggestion.impact_score
                    }
                    for suggestion in modification_result.suggestions
                ]
            }
        }
        
        return jsonify(result_data)
        
    except Exception as e:
        return jsonify({
            "error": f"خطأ في تعديل النص: {str(e)}"
        }), 500

@app.route('/api/text/preview-expansion', methods=['POST'])
def preview_text_expansion():
    """معاينة توسع النص قبل التطبيق"""
    if not TEXT_PROCESSING_ENABLED:
        return jsonify({"error": "خدمة معالجة النصوص غير متاحة"}), 503
    
    try:
        data = request.json
        text = data.get('text', '')
        target_length = data.get('target_length', 0)
        
        if not text:
            return jsonify({"error": "النص مطلوب"}), 400
        
        current_words = len(text.split())
        
        if target_length <= current_words:
            return jsonify({
                "success": True,
                "preview": {
                    "current_word_count": current_words,
                    "target_word_count": target_length,
                    "expansion_needed": False,
                    "message": "النص بالطول المطلوب أو أطول"
                }
            })
        
        expansion_ratio = target_length / current_words
        estimated_reading_time = target_length / 200  # كلمة في الدقيقة
        
        # تقدير نوع التوسع المطلوب
        expansion_type = "خفيف"
        if expansion_ratio > 1.5:
            expansion_type = "متوسط"
        if expansion_ratio > 2.0:
            expansion_type = "كبير"
        if expansion_ratio > 3.0:
            expansion_type = "كبير جداً"
        
        return jsonify({
            "success": True,
            "preview": {
                "current_word_count": current_words,
                "target_word_count": target_length,
                "words_to_add": target_length - current_words,
                "expansion_ratio": round(expansion_ratio, 2),
                "expansion_type": expansion_type,
                "estimated_reading_time": round(estimated_reading_time, 1),
                "expansion_needed": True,
                "recommendations": [
                    "إضافة تفاصيل وصفية" if expansion_ratio > 1.2 else None,
                    "إضافة سياق إضافي" if expansion_ratio > 1.5 else None,
                    "إضافة انتقالات طبيعية" if expansion_ratio > 1.3 else None,
                    "إضافة أمثلة وتوضيحات" if expansion_ratio > 2.0 else None
                ]
            }
        })
        
    except Exception as e:
        return jsonify({
            "error": f"خطأ في معاينة التوسع: {str(e)}"
        }), 500

@app.route('/api/text/preview-summarization', methods=['POST'])
def preview_text_summarization():
    """معاينة تلخيص النص قبل التطبيق"""
    if not TEXT_PROCESSING_ENABLED:
        return jsonify({"error": "خدمة معالجة النصوص غير متاحة"}), 503
    
    try:
        data = request.json
        text = data.get('text', '')
        target_length = data.get('target_length', 0)
        
        if not text:
            return jsonify({"error": "النص مطلوب"}), 400
        
        current_words = len(text.split())
        
        if target_length >= current_words:
            return jsonify({
                "success": True,
                "preview": {
                    "current_word_count": current_words,
                    "target_word_count": target_length,
                    "summarization_needed": False,
                    "message": "النص بالطول المطلوب أو أقصر"
                }
            })
        
        compression_ratio = target_length / current_words
        words_to_remove = current_words - target_length
        
        # تقدير نوع التلخيص المطلوب
        summarization_type = "خفيف"
        if compression_ratio < 0.7:
            summarization_type = "متوسط"
        if compression_ratio < 0.5:
            summarization_type = "كبير"
        if compression_ratio < 0.3:
            summarization_type = "كبير جداً"
        
        return jsonify({
            "success": True,
            "preview": {
                "current_word_count": current_words,
                "target_word_count": target_length,
                "words_to_remove": words_to_remove,
                "compression_ratio": round(compression_ratio, 2),
                "summarization_type": summarization_type,
                "summarization_needed": True,
                "warnings": [
                    "قد يؤثر على التفاصيل المهمة" if compression_ratio < 0.5 else None,
                    "تأكد من مراجعة النتيجة" if compression_ratio < 0.7 else None,
                    "قد يحتاج تعديل يدوي" if compression_ratio < 0.3 else None
                ]
            }
        })
        
    except Exception as e:
        return jsonify({
            "error": f"خطأ في معاينة التلخيص: {str(e)}"
        }), 500

# ==============================================================================
# APIs نظام الشاهد - Witness Content Processing APIs
# ==============================================================================

@app.route('/api/witness/process', methods=['POST'])
def process_witness_content():
    """معالجة وتحليل محتوى الشاهد"""
    if not TEXT_PROCESSING_ENABLED:
        return jsonify({"error": "خدمة معالجة النصوص غير متاحة"}), 503
    
    try:
        data = request.json
        content = data.get('content', '')
        content_type = data.get('content_type', 'TEXT_DOCUMENT')  # VIDEO_TRANSCRIPT, AUDIO_TRANSCRIPT, etc.
        source_url = data.get('source_url')
        
        if not content or len(content.strip()) < 20:
            return jsonify({"error": "محتوى الشاهد قصير جداً للمعالجة"}), 400
        
        # تحويل نوع المحتوى
        from advanced_text_processing_service import WitnessContentType
        
        content_type_enum = {
            'VIDEO_TRANSCRIPT': WitnessContentType.VIDEO_TRANSCRIPT,
            'AUDIO_TRANSCRIPT': WitnessContentType.AUDIO_TRANSCRIPT,
            'TEXT_DOCUMENT': WitnessContentType.TEXT_DOCUMENT,
            'PDF_DOCUMENT': WitnessContentType.PDF_DOCUMENT,
            'WEB_CONTENT': WitnessContentType.WEB_CONTENT
        }.get(content_type, WitnessContentType.TEXT_DOCUMENT)
        
        # تشغيل المعالجة
        loop = asyncio.new_event_loop()
        asyncio.set_event_loop(loop)
        witness_analysis = loop.run_until_complete(
            text_processor.process_witness_content(content, content_type_enum, source_url)
        )
        loop.close()
        
        # تحويل النتيجة لصيغة JSON
        result_data = {
            "success": True,
            "witness_analysis": {
                "witness_content": {
                    "id": witness_analysis.witness_content.id,
                    "content_type": witness_analysis.witness_content.content_type.value,
                    "source_url": witness_analysis.witness_content.source_url,
                    "title": witness_analysis.witness_content.title,
                    "content": witness_analysis.witness_content.content,
                    "timestamp": witness_analysis.witness_content.timestamp.isoformat(),
                    "metadata": witness_analysis.witness_content.metadata,
                    "credibility_score": witness_analysis.witness_content.credibility_score,
                    "relevance_score": witness_analysis.witness_content.relevance_score
                },
                "key_information": witness_analysis.key_information,
                "factual_claims": witness_analysis.factual_claims,
                "narrative_elements": witness_analysis.narrative_elements,
                "integration_suggestions": witness_analysis.integration_suggestions,
                "credibility_assessment": witness_analysis.credibility_assessment
            }
        }
        
        return jsonify(result_data)
        
    except Exception as e:
        return jsonify({
            "error": f"خطأ في معالجة محتوى الشاهد: {str(e)}"
        }), 500

@app.route('/api/witness/upload-transcript', methods=['POST'])
def upload_transcript():
    """رفع ترانسكريبت فيديو أو صوت"""
    try:
        # فحص إذا كان الملف موجود
        if 'file' not in request.files:
            return jsonify({"error": "لم يتم رفع أي ملف"}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({"error": "لم يتم اختيار ملف"}), 400
        
        # قراءة محتوى الملف
        try:
            content = file.read().decode('utf-8')
        except UnicodeDecodeError:
            try:
                content = file.read().decode('utf-8-sig')  # للملفات مع BOM
            except:
                return jsonify({"error": "تعذر قراءة الملف. تأكد أنه ملف نصي بترميز UTF-8"}), 400
        
        if not content or len(content.strip()) < 20:
            return jsonify({"error": "محتوى الملف قصير جداً أو فارغ"}), 400
        
        # تحديد نوع المحتوى من اسم الملف أو البيانات
        filename = file.filename.lower()
        if 'video' in filename or 'transcript' in filename:
            content_type = 'VIDEO_TRANSCRIPT'
        elif 'audio' in filename or 'speech' in filename:
            content_type = 'AUDIO_TRANSCRIPT'
        else:
            content_type = 'TEXT_DOCUMENT'
        
        # معالجة المحتوى
        if TEXT_PROCESSING_ENABLED:
            from advanced_text_processing_service import WitnessContentType
            
            content_type_enum = {
                'VIDEO_TRANSCRIPT': WitnessContentType.VIDEO_TRANSCRIPT,
                'AUDIO_TRANSCRIPT': WitnessContentType.AUDIO_TRANSCRIPT,
                'TEXT_DOCUMENT': WitnessContentType.TEXT_DOCUMENT
            }.get(content_type, WitnessContentType.TEXT_DOCUMENT)
            
            loop = asyncio.new_event_loop()
            asyncio.set_event_loop(loop)
            witness_analysis = loop.run_until_complete(
                text_processor.process_witness_content(content, content_type_enum)
            )
            loop.close()
            
            # إرجاع تحليل مختصر
            return jsonify({
                "success": True,
                "message": "تم رفع الترانسكريبت ومعالجته بنجاح",
                "file_info": {
                    "filename": file.filename,
                    "size": len(content),
                    "content_type": content_type
                },
                "analysis_summary": {
                    "title": witness_analysis.witness_content.title,
                    "credibility_score": witness_analysis.witness_content.credibility_score,
                    "key_information_count": len(witness_analysis.key_information),
                    "factual_claims_count": len(witness_analysis.factual_claims),
                    "narrative_elements_count": len(witness_analysis.narrative_elements)
                },
                "witness_id": witness_analysis.witness_content.id
            })
        else:
            # إرجاع معلومات أساسية فقط
            return jsonify({
                "success": True,
                "message": "تم رفع الترانسكريبت بنجاح",
                "file_info": {
                    "filename": file.filename,
                    "size": len(content),
                    "content_type": content_type
                },
                "content_preview": content[:200] + "..." if len(content) > 200 else content
            })
        
    except Exception as e:
        return jsonify({
            "error": f"خطأ في رفع الترانسكريبت: {str(e)}"
        }), 500

@app.route('/api/witness/integrate-suggestions', methods=['POST'])
def get_integration_suggestions():
    """الحصول على اقتراحات دمج محتوى الشاهد في النص"""
    if not TEXT_PROCESSING_ENABLED:
        return jsonify({"error": "خدمة معالجة النصوص غير متاحة"}), 503
    
    try:
        data = request.json
        witness_content = data.get('witness_content', '')
        current_text = data.get('current_text', '')
        integration_context = data.get('context', '')
        
        if not witness_content:
            return jsonify({"error": "محتوى الشاهد مطلوب"}), 400
        
        # توليد اقتراحات الدمج
        loop = asyncio.new_event_loop()
        asyncio.set_event_loop(loop)
        suggestions = loop.run_until_complete(
            text_processor._generate_integration_suggestions(witness_content)
        )
        loop.close()
        
        # تحليل النص الحالي لتحديد أفضل نقاط الدمج
        integration_points = []
        if current_text:
            sentences = current_text.split('.')
            for i, sentence in enumerate(sentences):
                if len(sentence.strip()) > 10:
                    # تحديد النقاط المناسبة للدمج
                    if any(keyword in sentence for keyword in ['الواقع', 'الحقيقة', 'المؤكد', 'بناءً على']):
                        integration_points.append({
                            'position': i,
                            'sentence': sentence.strip(),
                            'reason': 'مناسب لإدراج معلومات واقعية'
                        })
        
        return jsonify({
            "success": True,
            "integration_suggestions": suggestions,
            "integration_points": integration_points[:5],  # أهم 5 نقاط
            "usage_examples": [
                "ادمج المعلومات كحوار للشخصيات",
                "استخدمها كخلفية تاريخية للأحداث",
                "أضفها كتفاصيل وصفية للمكان أو الزمان",
                "استوحِ منها تطوير الشخصيات",
                "استخدمها كأساس للصراع في القصة"
            ]
        })
        
    except Exception as e:
        return jsonify({
            "error": f"خطأ في توليد اقتراحات الدمج: {str(e)}"
        }), 500

if __name__ == '__main__':
    print("🚀 بدء تشغيل خادم الذكاء الاصطناعي...")
    print("🔗 الرابط: http://localhost:5000")
    print("📚 المشاريع سيتم حفظها في: novel_projects.db")
    print("🔑 تأكد من تكوين مفاتيح API في ملف .env")
    print("⚙️  محرك سير العمل متاح الآن!")
    
    # إظهار حالة نظام الوكلاء المتقدم
    if AGENT_STUDIO_ENABLED:
        print("🤖 استوديو الوكلاء المتقدم متاح الآن!")
    else:
        print("⚠️  استوديو الوكلاء المتقدم غير متاح")

# ==============================================
# APIs المحرر التفاعلي المتقدم
# ==============================================

@app.route('/api/advanced-editor/analyze', methods=['POST'])
def analyze_text_advanced():
    """تحليل النص المتقدم للمحرر التفاعلي"""
    try:
        data = request.get_json()
        text = data.get('text', '').strip()
        
        if not text:
            return jsonify({'error': 'النص مطلوب'}), 400
        
        # حساب الإحصائيات الأساسية
        words = text.split()
        sentences = text.count('.') + text.count('!') + text.count('?')
        paragraphs = len([p for p in text.split('\n') if p.strip()])
        
        # تحليل سهولة القراءة (مبسط)
        avg_words_per_sentence = len(words) / max(sentences, 1)
        readability_score = max(0, min(1, (100 - avg_words_per_sentence * 2) / 100))
        
        # تحليل الجودة (مبسط)
        unique_words = len(set(words))
        vocabulary_diversity = unique_words / max(len(words), 1)
        quality_score = min(1, vocabulary_diversity * 1.5)
        
        # تحليل الأسلوب
        style_score = 0.7  # قيمة افتراضية
        
        # تحليل المشاعر (مبسط)
        emotions = {
            'إيجابي': 0.4,
            'سلبي': 0.2,
            'محايد': 0.4
        }
        
        # اكتشاف المشاكل
        issues = []
        if avg_words_per_sentence > 20:
            issues.append({
                'type': 'جمل طويلة',
                'text': 'الجمل طويلة جداً',
                'position': 0,
                'severity': 'medium',
                'suggestion': 'حاول تقسيم الجمل الطويلة إلى جمل أقصر'
            })
        
        # اقتراحات عامة
        suggestions = [
            'يمكن تحسين التنوع في المفردات',
            'فكر في إضافة المزيد من التفاصيل الحسية',
            'راجع ترابط الأفكار بين الفقرات'
        ]
        
        analysis = {
            'wordCount': len(words),
            'sentenceCount': sentences,
            'paragraphCount': paragraphs,
            'readabilityScore': readability_score,
            'qualityScore': quality_score,
            'styleScore': style_score,
            'emotions': emotions,
            'suggestions': suggestions,
            'issues': issues
        }
        
        return jsonify(analysis)
        
    except Exception as e:
        return jsonify({'error': f'خطأ في تحليل النص: {str(e)}'}), 500


@app.route('/api/advanced-editor/modify-length', methods=['POST'])
def modify_text_length():
    """تعديل طول النص (إطالة أو تقصير)"""
    try:
        data = request.get_json()
        text = data.get('text', '').strip()
        modification_type = data.get('type', 'expand')  # expand أو compress
        target_length = data.get('targetLength', 2000)
        preserve_style = data.get('preserveStyle', True)
        
        if not text:
            return jsonify({'error': 'النص مطلوب'}), 400
        
        current_words = len(text.split())
        
        if modification_type == 'expand':
            # توسيع النص
            if current_words >= target_length:
                return jsonify({
                    'modifiedText': text,
                    'message': 'النص بالطول المطلوب أو أطول',
                    'wordCount': current_words
                })
            
            # إضافة تفاصيل وتوسيعات (مبسط)
            sentences = text.split('.')
            expanded_sentences = []
            
            for sentence in sentences:
                if sentence.strip():
                    expanded_sentences.append(sentence.strip())
                    # إضافة تفاصيل إضافية
                    if len(expanded_sentences) % 2 == 0:
                        expanded_sentences.append("وفي هذا السياق، يمكن أن نضيف المزيد من التفاصيل والوصف الدقيق")
            
            modified_text = '. '.join(expanded_sentences)
            
        else:  # compress
            # ضغط النص
            if current_words <= target_length:
                return jsonify({
                    'modifiedText': text,
                    'message': 'النص بالطول المطلوب أو أقصر',
                    'wordCount': current_words
                })
            
            # تلخيص مبسط - أخذ النصف الأول
            words = text.split()
            modified_words = words[:target_length]
            modified_text = ' '.join(modified_words)
        
        new_word_count = len(modified_text.split())
        
        return jsonify({
            'modifiedText': modified_text,
            'wordCount': new_word_count,
            'originalWordCount': current_words,
            'modificationType': modification_type,
            'confidence': 0.85
        })
        
    except Exception as e:
        return jsonify({'error': f'خطأ في تعديل طول النص: {str(e)}'}), 500


@app.route('/api/advanced-editor/suggestions', methods=['POST'])
def get_smart_suggestions():
    """الحصول على اقتراحات ذكية للنص المحدد"""
    try:
        data = request.get_json()
        selected_text = data.get('text', '').strip()
        context = data.get('context', '')
        full_context = data.get('fullContext', False)
        
        if not selected_text:
            return jsonify({'error': 'النص المحدد مطلوب'}), 400
        
        # توليد اقتراحات ذكية
        suggestions = []
        
        # اقتراحات تحسين الأسلوب
        if len(selected_text.split()) > 10:
            suggestions.append("يمكن تحسين التدفق السردي في هذا المقطع")
            suggestions.append("فكر في إضافة المزيد من التفاصيل الحسية")
        
        # اقتراحات لغوية
        if '،' not in selected_text and len(selected_text.split()) > 5:
            suggestions.append("يمكن تحسين علامات الترقيم في هذه الجملة")
        
        # اقتراحات إبداعية
        suggestions.extend([
            "يمكن إضافة استعارة أو صورة بيانية هنا",
            "فكر في تنويع بداية الجملة",
            "يمكن ربط هذا المقطع بالسياق العام بشكل أفضل"
        ])
        
        return jsonify({
            'suggestions': suggestions[:4],  # أقصى 4 اقتراحات
            'confidence': 0.78,
            'suggestionType': 'style_improvement'
        })
        
    except Exception as e:
        return jsonify({'error': f'خطأ في توليد الاقتراحات: {str(e)}'}), 500


@app.route('/api/witness/upload', methods=['POST'])
def upload_witness_transcript():
    """رفع ترانسكريبت الشاهد"""
    try:
        if 'transcript' not in request.files:
            return jsonify({'error': 'ملف الترانسكريبت مطلوب'}), 400
        
        file = request.files['transcript']
        if file.filename == '':
            return jsonify({'error': 'لم يتم اختيار ملف'}), 400
        
        if not file.filename.endswith('.txt'):
            return jsonify({'error': 'يجب أن يكون الملف بصيغة .txt'}), 400
        
        # قراءة محتوى الملف
        content = file.read().decode('utf-8')
        
        # تحليل المحتوى واستخلاص المعلومات
        lines = content.split('\n')
        extracted_facts = []
        
        # استخلاص مبسط للحقائق
        for line in lines:
            if len(line.strip()) > 50 and any(keyword in line for keyword in ['قال', 'ذكر', 'أشار', 'أكد']):
                extracted_facts.append(line.strip()[:100])
        
        # تقييم المصداقية (مبسط)
        credibility_indicators = ['شاهد', 'رأيت', 'سمعت', 'كنت هناك']
        credibility_score = min(1.0, sum(1 for indicator in credibility_indicators if indicator in content) * 0.25)
        
        # اقتراحات الدمج
        suggested_integrations = [
            "يمكن استخدام هذه الشهادة لتعزيز المصداقية التاريخية",
            "هذه المعلومات يمكن دمجها في الوصف الخلفي للأحداث",
            "يمكن اقتباس أجزاء من هذه الشهادة مباشرة",
            "هذا المحتوى يوفر سياقاً ثقافياً مهماً"
        ]
        
        # إنشاء معرف فريد
        witness_id = f"witness_{int(time.time())}"
        
        return jsonify({
            'id': witness_id,
            'content': content,
            'extractedFacts': extracted_facts[:10],  # أقصى 10 حقائق
            'credibilityScore': credibility_score,
            'suggestedIntegrations': suggested_integrations,
            'wordCount': len(content.split()),
            'status': 'processed'
        })
        
    except Exception as e:
        return jsonify({'error': f'خطأ في رفع ملف الشاهد: {str(e)}'}), 500


@app.route('/api/witness/integrate', methods=['POST'])
def integrate_witness_info():
    """دمج معلومات من الشاهد في النص"""
    try:
        data = request.get_json()
        witness_id = data.get('witnessId')
        query = data.get('query', '')
        current_text = data.get('currentText', '')
        context = data.get('context', 'novel_chapter')
        
        if not witness_id:
            return jsonify({'error': 'معرف الشاهد مطلوب'}), 400
        
        # نص تجريبي للدمج (في التطبيق الحقيقي سيتم استرجاع البيانات من قاعدة البيانات)
        integrated_text = f"""
        
[بناءً على شهادة موثقة]: وفي هذا السياق، تشير الوثائق التاريخية إلى أن "{query}" كان له تأثير كبير على الأحداث. وكما ذكر أحد الشهود المعاصرين: "لقد شهدت بأم عيني كيف تطورت الأحداث في ذلك الوقت، وكان الجو مشحوناً بالتوتر والترقب."

هذه الشهادة توفر لنا نظرة عميقة على الحالة النفسية والاجتماعية السائدة آنذاك، مما يضفي على السرد بُعداً إضافياً من المصداقية والعمق التاريخي."""
        
        return jsonify({
            'integratedText': integrated_text.strip(),
            'integrationType': 'contextual_insertion',
            'confidenceScore': 0.82,
            'suggestedPosition': 'end_of_paragraph'
        })
        
    except Exception as e:
        return jsonify({'error': f'خطأ في دمج معلومات الشاهد: {str(e)}'}), 500


@app.route('/api/advanced-editor/rate', methods=['POST'])
def rate_content():
    """تقييم المحتوى للتعلم التكيفي"""
    try:
        data = request.get_json()
        text = data.get('text', '')
        rating = data.get('rating', 'neutral')  # positive, negative, neutral
        context = data.get('context', 'editing_session')
        
        # حفظ التقييم للتعلم التكيفي (هنا يمكن استخدام adaptive_learning_service)
        feedback_data = {
            'text_length': len(text.split()),
            'rating': rating,
            'context': context,
            'timestamp': datetime.now().isoformat(),
            'user_id': 'default_user'  # في التطبيق الحقيقي سيكون من نظام المصادقة
        }
        
        # في التطبيق الحقيقي سيتم حفظ هذا في قاعدة البيانات
        print(f"تم حفظ التقييم: {feedback_data}")
        
        return jsonify({
            'message': 'تم حفظ التقييم بنجاح',
            'rating': rating,
            'status': 'saved'
        })
        
    except Exception as e:
        return jsonify({'error': f'خطأ في حفظ التقييم: {str(e)}'}), 500

    
    # تشغيل الخادم
    app.run(
        debug=True,
        host='0.0.0.0',  # للسماح بالاتصال من الخارج
        port=5000,
        threaded=True  # لدعم طلبات متعددة
    )
